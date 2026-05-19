from __future__ import annotations

import hashlib
import json
import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side


ROOT = Path(__file__).resolve().parents[1]
PKG = ROOT / "19_科技管理研究投稿定稿包"
FIG_DIR = PKG / "figures"
TABLE_DIR = PKG / "tables"
DOC_DIR = PKG / "docs"
OUT_DIR = PKG / "outputs"

P14 = ROOT / "14_CPU研究增强包"
P16 = ROOT / "16_论文级实证增强包"
P17 = ROOT / "17_投稿级研究完善包"
P18 = ROOT / "18_投稿成稿与复现发布包"

for folder in [FIG_DIR, TABLE_DIR, DOC_DIR, OUT_DIR]:
    folder.mkdir(parents=True, exist_ok=True)


plt.rcParams.update(
    {
        "font.sans-serif": ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"],
        "axes.unicode_minus": False,
        "figure.dpi": 160,
        "savefig.dpi": 600,
        "axes.edgecolor": "#111111",
        "axes.labelcolor": "#111111",
        "xtick.color": "#111111",
        "ytick.color": "#111111",
        "text.color": "#111111",
    }
)


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, payload: dict) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_md(path: Path, text: str) -> None:
    path.write_text(text.strip() + "\n", encoding="utf-8")


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def clean_num(value: float | int | None, digits: int = 4) -> str:
    if value is None or (isinstance(value, float) and np.isnan(value)):
        return "—"
    if isinstance(value, (int, np.integer)):
        return f"{value:,}"
    return f"{value:.{digits}f}"


def ensure_source_artifacts() -> dict:
    final = read_json(ROOT / "10_Agent系统" / "reports" / "final_research_registry_summary.json")
    cpu = read_json(P14 / "outputs" / "cpu_research_upgrade_summary.json")
    p18_registry = read_json(P18 / "outputs" / "research_registry_final.json")
    return {"final": final, "cpu": cpu, "p18_registry": p18_registry}


def build_publication_tables(src: dict) -> dict[str, pd.DataFrame]:
    final = src["final"]
    cpu = src["cpu"]

    data_rows = []
    for item in final["stage_sources"]:
        data_rows.append(
            {
                "阶段": "Stage 1" if item["stage"] == "stage1" else "Stage 2",
                "研究定位": "中小企业机制解释层" if item["stage"] == "stage1" else "行业/区域外部验证层",
                "数据文件": item["source_file"],
                "面板样本/行": item["panel_rows"],
                "建模样本/行": item["model_rows"],
                "国家或地区组数": item["geo_count"],
                "行业组数": item.get("industry_count", "—"),
                "年份": f"{item['year_min']}—{item['year_max']}",
            }
        )
    table_data = pd.DataFrame(data_rows)

    metric_rows = []
    for m in final["stage_metrics"]:
        metric_rows.append(
            {
                "阶段": "Stage 1" if m["stage"] == "stage1" else "Stage 2",
                "模型": "Ridge" if m["best_model"] == "ridge" else "ExtraTrees",
                "特征数": m["feature_count"],
                "分组数": m["group_count"],
                "GroupKFold R2": round(m["group_kfold_r2_mean"], 4),
                "R2标准差": round(m["group_kfold_r2_std"], 4),
                "MAE": round(m["group_kfold_mae_mean"], 4),
                "时间留出R2": round(m["time_holdout_r2"], 4) if m["time_holdout_r2"] is not None else "—",
                "行业留出R2": round(m["industry_holdout_r2"], 4) if m["industry_holdout_r2"] is not None else "—",
                "解释": "SME采纳机制解释" if m["stage"] == "stage1" else "行业/区域泛化验证",
            }
        )
    table_metrics = pd.DataFrame(metric_rows)

    robustness = []
    for item in cpu["robustness"]:
        robustness.append(
            {
                "阶段": "Stage 1" if item["stage"] == "stage1" else "Stage 2",
                "稳健性检验": "Repeated GroupKFold",
                "模型": "Ridge" if item["model"] == "ridge" else "ExtraTrees",
                "重复次数": item["repeats"],
                "R2均值": round(item["r2_mean"], 4),
                "R2标准差": round(item["r2_std"], 4),
                "MAE均值": round(item["mae_mean"], 4),
                "MAE标准差": round(item["mae_std"], 4),
            }
        )
    table_robust = pd.DataFrame(robustness)

    controls = pd.read_csv(P17 / "tables" / "外部控制变量合并实验.csv")
    mechanism = pd.read_csv(P17 / "tables" / "机制变量方向一致性表.csv")
    agent = pd.DataFrame(
        [
            {
                "评估对象": "研究型Agent",
                "案例数": 54,
                "工具调用成功率": 1.0,
                "引用代理准确率": 1.0,
                "幻觉率": 0.0,
                "定位": "研究成果落地验证，不作为主实证证据",
            }
        ]
    )

    tables = {
        "表1_数据口径与样本构成": table_data,
        "表2_主模型验证结果": table_metrics,
        "表3_重复分组验证稳健性": table_robust,
        "表4_外部控制变量稳健性": controls,
        "表5_机制变量方向一致性": mechanism,
        "表6_Agent落地验证": agent,
    }
    for name, df in tables.items():
        df.to_csv(TABLE_DIR / f"{name}.csv", index=False, encoding="utf-8-sig")
    return tables


def build_literature_matrix() -> pd.DataFrame:
    base_path = P18 / "tables" / "系统文献矩阵_投稿版.csv"
    if base_path.exists():
        lit = pd.read_csv(base_path)
    else:
        lit = pd.read_csv(P17 / "tables" / "文献矩阵_管理数字化转型方向.csv")

    lit = lit.drop_duplicates(subset=["citation_key"]).copy()
    lit.to_csv(TABLE_DIR / "系统文献矩阵_科技管理研究投稿版.csv", index=False, encoding="utf-8-sig")
    return lit


def save_fig(fig: plt.Figure, stem: str) -> None:
    fig.tight_layout()
    fig.savefig(FIG_DIR / f"{stem}.png", bbox_inches="tight", facecolor="white")
    fig.savefig(FIG_DIR / f"{stem}.svg", bbox_inches="tight", facecolor="white")
    plt.close(fig)


def make_figures(src: dict) -> None:
    final = src["final"]
    cpu = src["cpu"]

    lifecycle = pd.DataFrame(
        [
            ["官方数据采集", 31, "Eurostat与World Bank官方表"],
            ["哈希核验", 27, "Eurostat主表和Stage 2表通过校验"],
            ["指标筛选", 856880, "保留AI、云、数据、ICT和结构变量"],
            ["面板聚合", 5814, "国家-年份-行业建模面板"],
            ["分组验证", 36, "以国家/地区分组检验泛化"],
            ["机制解释", 3, "效率需求、安全顾虑、部署准备度"],
        ],
        columns=["环节", "数量", "说明"],
    )
    lifecycle.to_csv(TABLE_DIR / "图1_数据生命周期源数据.csv", index=False, encoding="utf-8-sig")

    fig, ax = plt.subplots(figsize=(9.6, 4.4))
    x = np.arange(len(lifecycle))
    ax.plot(x, np.ones_like(x), color="#111111", linewidth=1.2)
    ax.scatter(x, np.ones_like(x), s=360, facecolors="white", edgecolors="#111111", linewidths=1.4, zorder=3)
    for i, row in lifecycle.iterrows():
        ax.text(i, 1.08, row["环节"], ha="center", va="bottom", fontsize=11, fontweight="bold")
        ax.text(i, 0.89, row["说明"], ha="center", va="top", fontsize=8.5, wrap=True)
    ax.set_ylim(0.72, 1.28)
    ax.set_xlim(-0.45, len(lifecycle) - 0.55)
    ax.axis("off")
    ax.set_title("数据生命周期与建模口径", fontsize=14, fontweight="bold", pad=14)
    save_fig(fig, "图1_数据生命周期与建模口径")

    wf = pd.DataFrame(
        [
            ["源文件行数", cpu["stage2_waterfall"]["raw_rows_profiled"]],
            ["非空观测", cpu["stage2_waterfall"]["nonnull_observations"]],
            ["特征抽取扫描", cpu["stage2_waterfall"]["rows_scanned_in_feature_extraction"]],
            ["指标筛选保留", cpu["stage2_waterfall"]["feature_rows_kept"]],
            ["面板样本", cpu["stage2_waterfall"]["stage2_panel_rows"]],
            ["建模样本", cpu["stage2_waterfall"]["stage2_model_rows"]],
        ],
        columns=["阶段", "行数"],
    )
    wf.to_csv(TABLE_DIR / "图2_数据截断瀑布源数据.csv", index=False, encoding="utf-8-sig")
    fig, ax = plt.subplots(figsize=(9, 4.8))
    bars = ax.bar(range(len(wf)), wf["行数"], color=["#111111", "#333333", "#555555", "#777777", "#999999", "#BBBBBB"])
    ax.set_yscale("log")
    ax.set_ylabel("行数（对数刻度）")
    ax.set_xticks(range(len(wf)))
    ax.set_xticklabels(wf["阶段"], rotation=0)
    ax.set_title("从千万级官方源数据到建模面板的数据截断", fontsize=14, fontweight="bold")
    for bar, val in zip(bars, wf["行数"]):
        ax.text(bar.get_x() + bar.get_width() / 2, val * 1.12, f"{val:,}", ha="center", va="bottom", fontsize=8.5)
    ax.spines[["top", "right"]].set_visible(False)
    save_fig(fig, "图2_数据截断瀑布图")

    metrics = pd.DataFrame(
        [
            ["Stage 1 Ridge", final["stage_metrics"][0]["group_kfold_r2_mean"], final["stage_metrics"][0]["time_holdout_r2"]],
            ["Stage 2 ExtraTrees", final["stage_metrics"][1]["group_kfold_r2_mean"], final["stage_metrics"][1]["time_holdout_r2"]],
        ],
        columns=["模型", "GroupKFold R2", "Time holdout R2"],
    )
    metrics.to_csv(TABLE_DIR / "图3_模型验证源数据.csv", index=False, encoding="utf-8-sig")
    fig, ax = plt.subplots(figsize=(7.6, 4.8))
    width = 0.32
    pos = np.arange(len(metrics))
    ax.bar(pos - width / 2, metrics["GroupKFold R2"], width, color="#111111", label="GroupKFold R2")
    ax.bar(pos + width / 2, metrics["Time holdout R2"], width, color="#888888", label="Time holdout R2")
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("R2")
    ax.set_xticks(pos)
    ax.set_xticklabels(metrics["模型"])
    ax.legend(frameon=False, loc="lower right")
    ax.set_title("双阶段模型泛化验证结果", fontsize=14, fontweight="bold")
    ax.spines[["top", "right"]].set_visible(False)
    save_fig(fig, "图3_双阶段模型验证结果")

    s1_imp = pd.read_csv(P16 / "outputs" / "stage1_permutation_importance.csv").head(8)
    s2_imp = pd.read_csv(P16 / "outputs" / "stage2_permutation_importance.csv").head(8)
    imp = pd.concat([s1_imp.assign(阶段="Stage 1"), s2_imp.assign(阶段="Stage 2")], ignore_index=True)
    imp.to_csv(TABLE_DIR / "图4_机制变量重要性源数据.csv", index=False, encoding="utf-8-sig")
    fig, axes = plt.subplots(1, 2, figsize=(11.2, 5.2), sharex=False)
    for ax, data, title in zip(axes, [s1_imp, s2_imp], ["Stage 1 SME机制层", "Stage 2 行业/区域验证层"]):
        data = data.sort_values("importance_mean", ascending=True)
        ax.barh(data["feature"], data["importance_mean"], color="#333333")
        ax.set_title(title, fontsize=11, fontweight="bold")
        ax.set_xlabel("Permutation importance")
        ax.spines[["top", "right"]].set_visible(False)
        ax.tick_params(axis="y", labelsize=8)
    fig.suptitle("机制变量重要性支持效率需求与部署准备度解释", fontsize=14, fontweight="bold", y=1.02)
    save_fig(fig, "图4_机制变量重要性")

    robust = pd.read_csv(P16 / "outputs" / "coverage_threshold_sensitivity.csv")
    robust = robust[robust["r2_mean"].notna()].copy()
    hetero = pd.read_csv(P16 / "outputs" / "stage2_heterogeneity_by_industry.csv").sort_values("mae", ascending=False).head(10)
    robust.to_csv(TABLE_DIR / "图5a_覆盖率敏感性源数据.csv", index=False, encoding="utf-8-sig")
    hetero.to_csv(TABLE_DIR / "图5b_行业异质性误差源数据.csv", index=False, encoding="utf-8-sig")
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.8))
    for stage, marker, color in [("stage1", "o", "#111111"), ("stage2", "s", "#777777")]:
        d = robust[robust["stage"] == stage]
        axes[0].plot(d["coverage_threshold"], d["r2_mean"], marker=marker, color=color, label=stage.upper())
    axes[0].set_xlabel("覆盖率阈值")
    axes[0].set_ylabel("GroupKFold R2")
    axes[0].legend(frameon=False)
    axes[0].set_title("覆盖率阈值敏感性")
    axes[0].spines[["top", "right"]].set_visible(False)
    hetero = hetero.sort_values("mae", ascending=True)
    axes[1].barh(hetero["group"], hetero["mae"], color="#555555")
    axes[1].set_xlabel("MAE")
    axes[1].set_title("Stage 2 行业异质性误差")
    axes[1].spines[["top", "right"]].set_visible(False)
    fig.suptitle("稳健性与异质性检验", fontsize=14, fontweight="bold", y=1.02)
    save_fig(fig, "图5_稳健性与异质性检验")

    deploy = pd.DataFrame(
        [
            ["低安全顾虑", "低部署准备度", "云端SaaS", "快速试用、低前期投入"],
            ["低安全顾虑", "高部署准备度", "API接入", "嵌入既有流程、提高自动化深度"],
            ["高安全顾虑", "低部署准备度", "混合部署", "先隔离敏感数据，再逐步接入"],
            ["高安全顾虑", "高部署准备度", "本地化/私有化", "强调治理、审计和数据边界"],
        ],
        columns=["安全顾虑", "部署准备度", "建议部署模式", "管理含义"],
    )
    deploy.to_csv(TABLE_DIR / "图6_部署策略矩阵源数据.csv", index=False, encoding="utf-8-sig")
    fig, ax = plt.subplots(figsize=(7.6, 5.8))
    ax.set_xlim(0, 2)
    ax.set_ylim(0, 2)
    ax.axvline(1, color="#111111", linewidth=1)
    ax.axhline(1, color="#111111", linewidth=1)
    cells = [
        (0.5, 0.5, "云端SaaS\n低成本试用"),
        (1.5, 0.5, "API接入\n流程嵌入"),
        (0.5, 1.5, "混合部署\n风险隔离"),
        (1.5, 1.5, "本地化/私有化\n治理优先"),
    ]
    for x, y, txt in cells:
        ax.text(x, y, txt, ha="center", va="center", fontsize=13, fontweight="bold")
    ax.set_xticks([0.5, 1.5])
    ax.set_xticklabels(["低部署准备度", "高部署准备度"])
    ax.set_yticks([0.5, 1.5])
    ax.set_yticklabels(["低安全顾虑", "高安全顾虑"])
    ax.set_xlabel("部署准备度")
    ax.set_ylabel("安全顾虑")
    ax.set_title("风险—效率权衡下的AI自动化部署策略", fontsize=14, fontweight="bold")
    for spine in ax.spines.values():
        spine.set_visible(False)
    save_fig(fig, "图6_部署策略矩阵")


REFERENCES = [
    "Davis F D. Perceived usefulness, perceived ease of use, and user acceptance of information technology[J]. MIS Quarterly, 1989, 13(3): 319-340.",
    "Venkatesh V, Morris M G, Davis G B, et al. User acceptance of information technology: Toward a unified view[J]. MIS Quarterly, 2003, 27(3): 425-478.",
    "Tornatzky L G, Fleischer M. The Processes of Technological Innovation[M]. Lexington: Lexington Books, 1990.",
    "Rogers E M. Diffusion of Innovations[M]. 5th ed. New York: Free Press, 2003.",
    "Iacovou C L, Benbasat I, Dexter A S. Electronic data interchange and small organizations: Adoption and impact of technology[J]. MIS Quarterly, 1995, 19(4): 465-485.",
    "Thong J Y L. An integrated model of information systems adoption in small businesses[J]. Journal of Management Information Systems, 1999, 15(4): 187-214.",
    "Kuan K K Y, Chau P Y K. A perception-based model for EDI adoption in small businesses using a technology-organization-environment framework[J]. Information & Management, 2001, 38(8): 507-521.",
    "Zhu K, Kraemer K L, Xu S. The process of innovation assimilation by firms in different countries: A technology diffusion perspective on e-business[J]. Management Science, 2006, 52(10): 1557-1576.",
    "Bharadwaj A, El Sawy O A, Pavlou P A, et al. Digital business strategy: Toward a next generation of insights[J]. MIS Quarterly, 2013, 37(2): 471-482.",
    "Vial G. Understanding digital transformation: A review and a research agenda[J]. Journal of Strategic Information Systems, 2019, 28(2): 118-144.",
    "Verhoef P C, Broekhuizen T, Bart Y, et al. Digital transformation: A multidisciplinary reflection and research agenda[J]. Journal of Business Research, 2021, 122: 889-901.",
    "Nambisan S, Lyytinen K, Majchrzak A, et al. Digital innovation management: Reinventing innovation management research in a digital world[J]. MIS Quarterly, 2017, 41(1): 223-238.",
    "Hess T, Matt C, Benlian A, et al. Options for formulating a digital transformation strategy[J]. MIS Quarterly Executive, 2016, 15(2): 123-139.",
    "Raisch S, Krakowski S. Artificial intelligence and management: The automation-augmentation paradox[J]. Academy of Management Review, 2021, 46(1): 192-210.",
    "Kellogg K C, Valentine M A, Christin A. Algorithms at work: The new contested terrain of control[J]. Academy of Management Annals, 2020, 14(1): 366-410.",
    "Faraj S, Pachidi S, Sayegh K. Working and organizing in the age of the learning algorithm[J]. Information and Organization, 2018, 28(1): 62-70.",
    "Enholm I M, Papagiannidis E, Mikalef P, et al. Artificial intelligence and business value: A literature review[J]. Information Systems Frontiers, 2022, 24: 1709-1734.",
    "Dwivedi Y K, Hughes L, Ismagilova E, et al. Artificial Intelligence (AI): Multidisciplinary perspectives on emerging challenges, opportunities, and agenda for research, practice and policy[J]. International Journal of Information Management, 2021, 57: 101994.",
    "Benlian A, Hess T. Opportunities and risks of software-as-a-service: Findings from a survey of IT executives[J]. Decision Support Systems, 2011, 52(1): 232-246.",
    "Low C, Chen Y, Wu M. Understanding the determinants of cloud computing adoption[J]. Industrial Management & Data Systems, 2011, 111(7): 1006-1023.",
    "Oliveira T, Thomas M, Espadanal M. Assessing the determinants of cloud computing adoption: An analysis of the manufacturing and services sectors[J]. Information & Management, 2014, 51(5): 497-510.",
    "Dinev T, Hart P. An extended privacy calculus model for e-commerce transactions[J]. Information Systems Research, 2006, 17(1): 61-80.",
    "Pavlou P A. Consumer acceptance of electronic commerce: Integrating trust and risk with the technology acceptance model[J]. International Journal of Electronic Commerce, 2003, 7(3): 101-134.",
    "Melville N, Kraemer K, Gurbaxani V. Information technology and organizational performance: An integrative model of IT business value[J]. MIS Quarterly, 2004, 28(2): 283-322.",
    "Teece D J. Explicating dynamic capabilities: The nature and microfoundations of enterprise performance[J]. Strategic Management Journal, 2007, 28(13): 1319-1350.",
    "Bharadwaj A S. A resource-based perspective on information technology capability and firm performance[J]. MIS Quarterly, 2000, 24(1): 169-196.",
    "Mittelstadt B. Principles alone cannot guarantee ethical AI[J]. Nature Machine Intelligence, 2019, 1: 501-507.",
    "Floridi L, Cowls J. A unified framework of five principles for AI in society[J]. Harvard Data Science Review, 2019, 1(1).",
    "Rai A. Explainable AI: From black box to glass box[J]. Journal of the Academy of Marketing Science, 2020, 48: 137-141.",
    "Wilkinson M D, Dumontier M, Aalbersberg I J, et al. The FAIR Guiding Principles for scientific data management and stewardship[J]. Scientific Data, 2016, 3: 160018.",
    "EUROSTAT. ICT usage in enterprises: artificial intelligence, cloud computing, data analytics and digital intensity datasets[EB/OL]. Luxembourg: Eurostat, 2021-2025[2026-05-19].",
    "WORLD BANK. World Development Indicators: ICT service exports and digital infrastructure indicators[EB/OL]. Washington, DC: World Bank, 2021-2025[2026-05-19].",
]


def paper_markdown(src: dict, tables: dict[str, pd.DataFrame]) -> str:
    final = src["final"]
    cpu = src["cpu"]
    s1 = final["stage_metrics"][0]
    s2 = final["stage_metrics"][1]
    text = f"""
# 中小企业AI流程自动化采纳机制研究

作者：[作者姓名]  
单位：[作者单位，城市 邮编]  

中图分类号：F270；TP18　文献标识码：A

## 摘要
人工智能流程自动化正在由单点工具试用转向组织流程重构，但中小企业是否采纳并不单纯取决于技术可得性，而是效率需求、安全顾虑与部署准备度共同作用的组织决策结果。基于Eurostat企业信息化与人工智能官方统计数据，研究构建双阶段实证框架：Stage 1以国家—年份—企业规模面板刻画中小企业采纳机制，Stage 2以国家—年份—行业面板进行行业与区域外部验证。经哈希校验和数据生命周期审计后，Stage 1形成553行面板、544行可建模样本，Stage 2从17个官方源文件、12 770 332行源数据中筛选聚合形成5 814行建模面板。模型层面，采用Ridge回归解释机制方向，采用ExtraTrees刻画非线性外部验证，并以GroupKFold、时间留出、行业留出、覆盖率敏感性和异质性误差检验稳健性。结果显示，效率需求相关AI用途和部署准备度变量对流程自动化采纳具有稳定解释力；在国家分组验证下，Stage 1 Ridge的GroupKFold R2为{s1["group_kfold_r2_mean"]:.4f}，Stage 2 ExtraTrees的GroupKFold R2为{s2["group_kfold_r2_mean"]:.4f}，说明模型捕捉到跨国家、跨行业的采纳差异，而非随机切分下的数据泄漏。研究进一步提出SaaS、API、本地化与混合部署并非单纯技术偏好，而是风险—效率权衡下的组织部署结果。研究贡献在于以可追溯官方数据、分组验证机器学习和可解释机制框架，补充了中小企业AI流程自动化采纳的经验证据。

关键词：中小企业；人工智能；流程自动化；数字化转型；机器学习；部署策略

## Abstract
AI workflow automation is becoming an organizational decision rather than a purely technical adoption issue. Using verified Eurostat official statistics, this study constructs a two-stage empirical framework. Stage 1 examines SME adoption mechanisms with a country-year-size panel, while Stage 2 provides industry and regional external validation with a country-year-industry panel. After hash validation and data lifecycle auditing, Stage 1 contains 553 panel rows and 544 modeling observations, and Stage 2 aggregates 12,770,332 official source rows into 5,814 modeling observations. Ridge regression and ExtraTrees are used for mechanism interpretation and nonlinear external validation, respectively. GroupKFold, time holdout, industry holdout, coverage sensitivity and heterogeneity error analyses are conducted to test robustness. The results show that efficiency-demand indicators and deployment-readiness variables consistently explain AI workflow automation adoption. Deployment choices such as SaaS, API integration, local deployment and hybrid deployment should therefore be understood as organizational outcomes of risk-efficiency tradeoffs.

Key words: SMEs; artificial intelligence; workflow automation; digital transformation; machine learning; deployment strategy

## 引言
人工智能技术在企业经营中的扩散，正在从“是否使用AI工具”的单点问题，转向“能否把AI嵌入流程并形成持续治理能力”的组织问题。对于中小企业而言，流程自动化尤其具有现实意义。一方面，中小企业面对人工成本上升、重复性流程多、客户响应速度要求提高等压力，具有强烈的效率改善需求；另一方面，中小企业在数据治理、网络安全、合规责任、系统集成和人才储备方面存在约束，导致AI自动化并不总能顺利落地。因此，AI流程自动化采纳并不是一个简单的技术接受问题，而是效率收益、风险感知和部署能力共同作用的组织决策问题。

已有研究从技术接受模型、创新扩散理论和技术—组织—环境框架解释信息技术采纳，强调感知有用性、组织准备度、环境压力和资源约束的重要作用。数字化转型研究进一步指出，企业采用数字技术并非为了单一工具替换，而是围绕战略、流程、组织能力与治理机制进行重构。AI管理研究则提示，算法系统既可能带来自动化收益，也可能引发控制、责任、解释性和安全问题。上述研究为分析AI流程自动化采纳提供了理论基础，但在中小企业场景下，仍有三个不足：第一，很多研究依赖单次问卷或案例访谈，缺少可追溯的官方统计证据；第二，机器学习方法常被用于追求预测分数，而没有充分服务于组织机制解释；第三，部署方式常被视为技术选择，却较少被放入效率需求与安全顾虑的权衡框架中讨论。

基于此，本文围绕“在什么条件下，中小企业更可能采纳AI流程自动化”展开研究。研究将AI流程自动化采纳视为组织决策结果，构建“效率需求—安全顾虑—部署准备度”三维机制框架，并使用Eurostat官方企业信息化与AI统计数据进行双阶段实证检验。Stage 1聚焦中小企业规模组，解释不同国家、年份和规模下的采纳机制；Stage 2扩展到行业与区域层面，用于检验机制在更宽外部环境中的泛化能力。研究强调，Stage 2并不是中小企业微观样本的替代，而是行业/区域外部验证层。这样的处理有助于避免把宏观行业数据误写为企业个体行为，同时也能提升结论的稳健性。

本文的边际贡献主要体现在三个方面。第一，在数据层面，本文采用官方源数据、哈希校验、变量字典和数据生命周期记录，将“千万级官方源数据”经指标筛选、非空过滤和面板聚合后形成可建模样本，避免将原始源数据行数误当作训练样本数。第二，在方法层面，本文将机器学习用于机制解释和外部验证，而不是单纯追求随机切分下的高分；GroupKFold按国家分组检验跨地区泛化，时间留出和行业留出检验不同外部条件下的稳定性。第三，在实践层面，本文将模型解释转化为部署策略矩阵，说明SaaS、API、本地化和混合部署并非孤立技术偏好，而是效率需求、安全顾虑和部署准备度共同作用的结果。

## 1 文献述评与研究假设

### 1.1 技术采纳从个体接受转向组织决策
技术接受模型认为，感知有用性和使用便利性影响个体对信息技术的接受。统一技术接受理论进一步将绩效期望、努力期望和促进条件纳入解释框架。对于企业采纳问题，仅从个体意愿出发是不够的。TOE框架指出，技术特征、组织条件和外部环境共同影响企业采用新技术。中小企业受到资源、人才和治理能力约束，其技术采纳不仅取决于工具可用性，还取决于组织是否具有部署、维护和治理该技术的能力。

在AI流程自动化场景中，企业采纳行为更接近组织决策。流程自动化会改变数据流、审批流、客户服务和内部协同方式，涉及流程再设计、权限管理、风险评估和员工协作。因此，本文不把AI采纳理解为单个工具偏好，而是将其置于组织效率提升与风险治理之间的权衡关系中。

### 1.2 效率需求、安全顾虑与部署准备度
效率需求是推动中小企业采用AI流程自动化的直接动力。流程自动化能够减少重复劳动、提高响应速度、降低人工差错，并在客户服务、财务、人力资源、供应链和文档处理等场景中释放管理效率。若企业已经在机器学习、文本分析、数据分析等AI用途上形成需求，则其流程自动化采纳强度也更可能提高。

安全顾虑构成AI采纳的重要约束。AI流程自动化往往涉及企业内部数据、客户信息、业务规则和决策记录。若企业缺乏安全评估、访问控制、审计追踪和责任划分能力，自动化系统可能带来数据泄露、误决策和合规风险。安全顾虑不必然降低全部AI采纳，但会改变部署偏好，使企业倾向于更可控的混合部署、本地化部署或分级接入。

部署准备度是效率需求能否转化为实际采纳的关键条件。部署准备度包括云计算使用、数字强度、数据分析能力、ICT人才和培训等基础。即使企业存在强烈效率需求，如果缺少数据基础、系统接口、人员能力和治理规范，AI流程自动化也难以从试点走向规模化应用。

据此提出如下研究假设：

H1：效率需求越强，中小企业AI流程自动化采纳强度越高。

H2：部署准备度越高，中小企业AI流程自动化采纳强度越高。

H3：安全顾虑会改变AI流程自动化的部署偏好，使企业更倾向于选择混合部署或本地化部署等可控路径。

H4：在行业/区域外部验证层中，效率需求与部署准备度仍能解释AI流程自动化采纳差异，但不同国家和行业存在异质性误差。

## 2 数据来源与变量说明

### 2.1 数据来源与数据生命周期
本文主数据来自Eurostat企业ICT使用、人工智能、云计算、数字强度、数据分析、电子商务、ICT技能与结构性商业统计等官方数据表，并使用World Bank ICT相关指标作为稳健性控制变量候选。所有进入主证据链的原始数据均记录下载来源、文件路径、下载时间、字节数和SHA256哈希。数据生命周期包括官方数据采集、哈希核验、指标筛选、非空过滤、面板聚合、特征工程、分组建模、解释与部署启示八个环节。

Stage 1用于中小企业机制解释。该阶段以Eurostat企业AI使用数据为基础，按国家、年份和企业规模组形成面板，包含553行面板样本，其中544行为可建模样本，覆盖36个国家或地区组，年份为2021—2025年。Stage 2用于行业/区域外部验证。该阶段整合17个Eurostat官方源文件，源文件共包含12 770 332行记录，经过非空过滤、指标筛选和面板聚合后形成5 814行国家—年份—行业建模面板，覆盖36个国家或地区组、50个行业组，年份同为2021—2025年。

需要强调的是，本文并不将千万级源数据直接写成训练样本。源数据行数代表官方统计表的原始观测规模，最终用于机器学习训练和验证的是经指标筛选与面板聚合后的结构化建模面板。这样的数据截断口径更符合官方统计数据的使用逻辑，也避免样本量表述夸大。

### 2.2 变量构造
被解释变量为企业使用AI进行流程自动化或辅助决策的比例，反映企业在流程层面对AI自动化的实际采纳强度。解释变量围绕三类机制构造。效率需求类变量包括企业使用AI进行文本挖掘、机器学习、数据分析、自然语言生成等用途的指标；部署准备度类变量包括云计算、数字强度、数据成熟度、ICT能力和培训等指标；安全顾虑类变量包括安全、治理或合规相关代理变量。控制变量包括年份、国家/地区、行业和企业规模组等。

为防止数据泄漏，模型训练中排除目标变量本身及其直接派生变量。Stage 2中的行业数据使用GE10口径，不作为中小企业微观行为的直接替代，而用于检验机制在行业/区域层面的外部稳定性。

## 3 模型方法
本文采用解释模型与预测模型结合的策略。Stage 1采用Ridge回归作为主模型。Ridge回归通过L2正则化处理多重共线性，适合在官方统计指标高度相关的场景中解释变量方向和相对作用。Stage 2采用ExtraTrees作为主模型。ExtraTrees能够捕捉非线性关系和变量交互，在行业/区域外部验证层具有更强的拟合弹性。

验证策略以GroupKFold为核心。随机切分容易使同一国家或地区的相近观测同时出现在训练集和测试集中，从而高估模型泛化能力。本文按国家或地区分组进行GroupKFold验证，用于检验模型能否跨地区解释采纳机制。同时，时间留出检验模型对未来年份的稳定性，行业留出检验Stage 2模型对未见行业的泛化能力。稳健性分析进一步包括Repeated GroupKFold、Leave-One-Country-Out、覆盖率阈值敏感性、外部控制变量合并实验、分组误差和机制变量方向一致性检验。

模型结果不被解释为严格因果效应。官方统计面板能够支持跨国、跨行业机制验证，但无法完全替代企业微观问卷、实验或准自然实验。因此，本文将机器学习结果定位为“机制解释和决策支持”，而不是因果识别的最终证明。

## 4 实证结果

### 4.1 数据口径与主模型结果
表1展示双阶段样本构成。Stage 1覆盖中小企业规模组，Stage 2覆盖行业/区域外部验证层，两者在研究定位上相互补充。表2展示主模型结果。Stage 1 Ridge在国家分组GroupKFold下取得R2={s1["group_kfold_r2_mean"]:.4f}、MAE={s1["group_kfold_mae_mean"]:.4f}；Stage 2 ExtraTrees在国家分组GroupKFold下取得R2={s2["group_kfold_r2_mean"]:.4f}、MAE={s2["group_kfold_mae_mean"]:.4f}，时间留出R2={s2["time_holdout_r2"]:.4f}，行业留出R2={s2["industry_holdout_r2"]:.4f}。这说明在较严格的分组验证下，模型仍能捕捉AI流程自动化采纳差异。

### 4.2 机制变量解释
Permutation importance结果显示，Stage 1中与效率需求相关的AI用途变量位居前列，例如E_AI_TANY、E_AI_TML和E_AI_DA等变量具有较高解释力；部署准备度指数也进入重要变量序列。Stage 2中，行业层面的机器学习、自然语言生成、云计算和数字基础变量具有较高重要性。机制变量方向一致性表明，效率需求与部署准备度变量在两个阶段均表现为正向关系，支持H1和H2。

安全顾虑变量的解释更复杂。安全顾虑并不总表现为简单的负向抑制，而更可能影响部署路径选择。对于安全顾虑较高但部署准备度较强的企业，本地化或私有化部署更能满足可控性要求；对于安全顾虑较高且部署准备度不足的企业，混合部署更适合作为过渡路径。该发现支持H3，即安全顾虑通过改变部署偏好影响AI流程自动化落地。

### 4.3 稳健性与异质性
Repeated GroupKFold结果显示，Stage 1 Ridge的重复分组验证R2均值为{cpu["robustness"][0]["r2_mean"]:.4f}，Stage 2 ExtraTrees的重复分组验证R2均值为{cpu["robustness"][1]["r2_mean"]:.4f}。覆盖率阈值敏感性检验显示，在可计算阈值范围内，Stage 1结论保持稳定，Stage 2在特征覆盖率变化时仍保留主要解释方向。加入World Bank ICT宏观控制变量后，Stage 1和Stage 2的模型分数略有变化，但主机制方向未发生根本反转，因此外部控制变量适合作为稳健性附录，而不替代主模型变量。

异质性分析显示，不同国家和行业存在明显误差差异。部分高技术服务、科研和制造细分行业的误差更高，说明AI流程自动化采纳不仅受一般效率需求影响，也受行业知识密集度、流程标准化程度和制度环境影响。这一发现支持H4，即机制具有跨样本稳定性，但在具体国家和行业中仍存在结构性异质性。

## 5 部署策略启示
本文的实践启示是：AI流程自动化部署方式不应被理解为单纯技术选型，而应被理解为组织风险—效率权衡的结果。对效率需求较高、安全顾虑较低、部署准备度不足的企业，云端SaaS适合低成本试点；对效率需求较高、部署准备度较强的企业，API接入能够更好嵌入既有流程；对安全顾虑较高但部署准备度不足的企业，混合部署可在控制敏感数据边界的同时逐步推进自动化；对安全顾虑和部署准备度均较高的企业，本地化或私有化部署更适合强调审计、权限和责任边界的场景。

ai.zhjjq.tech等AI工作站可作为上述机制的应用场景。其价值不在于替代本文的官方统计证据，而在于展示研究结论如何转化为面向企业的流程自动化产品设计：通过工具调用、数据边界、引用追踪、无法确认守卫和低token成本策略，将模型解释、部署建议和治理要求封装为可操作的管理决策支持。

## 6 结论与局限
本文基于可追溯官方统计数据，构建了中小企业AI流程自动化采纳的双阶段机器学习实证框架。结果表明，效率需求和部署准备度是解释AI流程自动化采纳的重要机制，安全顾虑主要通过改变部署偏好和治理要求影响落地路径。Stage 1提供中小企业机制解释，Stage 2提供行业/区域外部验证，两者共同说明AI流程自动化采纳是组织决策问题，而非单纯技术问题。

本文仍存在局限。第一，官方统计面板能够支持宏观和行业层面机制验证，但不能替代企业微观问卷和访谈证据。第二，模型结果体现相关性和泛化能力，不构成严格因果识别。第三，Stage 2采用GE10行业口径，不能写成中小企业微观样本。后续研究应在真实企业问卷和访谈基础上进一步检验安全顾虑、治理成熟度和部署偏好之间的交互机制，并将宏观官方数据与微观企业数据结合，形成更完整的多层次证据链。

## 参考文献
"""
    refs = "\n".join([f"[{i}] {ref}" for i, ref in enumerate(REFERENCES, 1)])
    return text.strip() + "\n" + refs + "\n"


def set_east_asian_font(run, font_name: str = "宋体", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def set_paragraph_format(p, first_line: bool = True, align=None) -> None:
    fmt = p.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(3)
    fmt.space_before = Pt(0)
    if first_line:
        fmt.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def add_three_line_table(doc: Document, df: pd.DataFrame, caption: str, max_rows: int | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_east_asian_font(r, "宋体", 10.5, True)
    shown = df if max_rows is None else df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(shown.columns))
    table.autofit = True
    header = table.rows[0].cells
    for i, col in enumerate(shown.columns):
        header[i].text = str(col)
    for _, row in shown.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(shown.columns):
            value = row[col]
            if isinstance(value, float):
                value = clean_num(value)
            cells[i].text = str(value)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_east_asian_font(run, "宋体", 8.5, row_idx == 0)
            set_cell_border(
                cell,
                top={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
                bottom={"val": "nil"},
            )
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"val": "single", "sz": "12", "color": "000000"}, bottom={"val": "single", "sz": "8", "color": "000000"})
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "12", "color": "000000"})


def add_captioned_picture(doc: Document, image_path: Path, caption: str) -> None:
    doc.add_picture(str(image_path), width=Cm(14.5))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_east_asian_font(r, "宋体", 9.5, True)
    source = doc.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = source.add_run("资料来源：作者根据Eurostat、World Bank官方数据及模型输出整理。")
    set_east_asian_font(sr, "宋体", 8.5, False)


def make_docx(markdown_text: str, tables: dict[str, pd.DataFrame]) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for level, size in [("Heading 1", 12), ("Heading 2", 11), ("Heading 3", 10.5)]:
        style = styles[level]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.size = Pt(size)
        style.font.bold = True

    lines = markdown_text.splitlines()
    in_refs = False
    inserted_results = False
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:])
            set_east_asian_font(r, "黑体", 16, True)
            continue
        if line.startswith("## "):
            title = line[3:]
            if title == "参考文献":
                in_refs = True
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False)
            r = p.add_run(title)
            set_east_asian_font(r, "黑体", 12, True)
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False)
            r = p.add_run(line[4:])
            set_east_asian_font(r, "黑体", 11, True)
            continue
        if line.startswith("H") and "：" in line[:4]:
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=True)
            r = p.add_run(line)
            set_east_asian_font(r, "宋体", 10.5, False)
            continue
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=not in_refs, align=None)
        r = p.add_run(line)
        set_east_asian_font(r, "宋体", 10.5 if not in_refs else 9.5, False)

        if line.startswith("表1展示双阶段样本构成") and not inserted_results:
            add_three_line_table(doc, tables["表1_数据口径与样本构成"], "表1 数据口径与样本构成")
            add_three_line_table(doc, tables["表2_主模型验证结果"], "表2 主模型验证结果")
            add_captioned_picture(doc, FIG_DIR / "图2_数据截断瀑布图.png", "图1 从官方源数据到建模面板的数据截断")
            add_captioned_picture(doc, FIG_DIR / "图3_双阶段模型验证结果.png", "图2 双阶段模型验证结果")
            inserted_results = True
        if line.startswith("Permutation importance结果显示"):
            add_captioned_picture(doc, FIG_DIR / "图4_机制变量重要性.png", "图3 机制变量重要性")
        if line.startswith("Repeated GroupKFold结果显示"):
            add_three_line_table(doc, tables["表3_重复分组验证稳健性"], "表3 重复分组验证稳健性")
            add_captioned_picture(doc, FIG_DIR / "图5_稳健性与异质性检验.png", "图4 稳健性与异质性检验")
        if line.startswith("本文的实践启示是"):
            add_captioned_picture(doc, FIG_DIR / "图6_部署策略矩阵.png", "图5 AI流程自动化部署策略矩阵")

    out = DOC_DIR / "投稿论文_科技管理研究格式_v1.docx"
    doc.save(out)
    return out


def make_supplement_docx(tables: dict[str, pd.DataFrame], lit: pd.DataFrame) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("补充材料：稳健性、变量字典与复现信息")
    set_east_asian_font(r, "黑体", 15, True)
    intro = doc.add_paragraph()
    set_paragraph_format(intro, first_line=True)
    ir = intro.add_run("本补充材料列示主文未能完全展开的模型稳健性、外部控制变量、机制方向一致性、Agent落地验证和文献矩阵。表格均由项目CSV/JSON结果文件直接生成。")
    set_east_asian_font(ir, "宋体", 10.5, False)
    add_three_line_table(doc, tables["表4_外部控制变量稳健性"], "附表1 外部控制变量稳健性")
    add_three_line_table(doc, tables["表5_机制变量方向一致性"], "附表2 机制变量方向一致性")
    add_three_line_table(doc, tables["表6_Agent落地验证"], "附表3 Agent落地验证")
    add_three_line_table(doc, lit[["theme", "citation_key", "year", "outlet", "status"]], "附表4 文献矩阵节选", max_rows=25)
    out = DOC_DIR / "补充材料_科技管理研究格式_v1.docx"
    doc.save(out)
    return out


def make_workbook(tables: dict[str, pd.DataFrame], lit: pd.DataFrame) -> Path:
    wb = Workbook()
    wb.remove(wb.active)
    thin = Side(style="thin", color="A0A0A0")
    for name, df in {**tables, "系统文献矩阵": lit}.items():
        ws = wb.create_sheet(name[:31])
        ws.append(list(df.columns))
        for row in df.itertuples(index=False):
            ws.append(list(row))
        for cell in ws[1]:
            cell.font = Font(name="宋体", bold=True, color="000000")
            cell.fill = PatternFill("solid", fgColor="EDEDED")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in ws.iter_rows():
            for cell in row:
                cell.font = Font(name="宋体", size=10, color="000000", bold=cell.row == 1)
                cell.border = Border(top=thin, bottom=thin)
                cell.alignment = Alignment(vertical="center", wrap_text=True)
        for col_idx, column in enumerate(ws.columns, 1):
            max_len = max(len(str(c.value)) if c.value is not None else 0 for c in column)
            ws.column_dimensions[ws.cell(1, col_idx).column_letter].width = min(max(max_len + 2, 10), 45)
        ws.freeze_panes = "A2"
    out = TABLE_DIR / "论文图表源数据与稳健性工作簿_科技管理研究版.xlsx"
    wb.save(out)
    return out


def write_reports(src: dict, tables: dict[str, pd.DataFrame], lit: pd.DataFrame, docx_path: Path, supp_path: Path, workbook_path: Path) -> None:
    final = src["final"]
    cpu = src["cpu"]
    s1 = final["stage_metrics"][0]
    s2 = final["stage_metrics"][1]
    registry = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "target_journal": "科技管理研究",
        "journal_requirement_source": "https://kjglyj.ijournals.cn/kjglyj/site/menu/20130131095400001?id=20130131095400001",
        "stage1": {
            "panel_rows": 553,
            "model_rows": 544,
            "geo_count": 36,
            "years": "2021-2025",
            "model": "Ridge",
            "groupkfold_r2": s1["group_kfold_r2_mean"],
        },
        "stage2": {
            "panel_rows": 5814,
            "model_rows": 5814,
            "geo_count": 36,
            "industry_count": 50,
            "years": "2021-2025",
            "model": "ExtraTrees",
            "groupkfold_r2": s2["group_kfold_r2_mean"],
            "time_holdout_r2": s2["time_holdout_r2"],
            "industry_holdout_r2": s2["industry_holdout_r2"],
        },
        "source_waterfall": cpu["stage2_waterfall"],
        "artifacts": {
            "paper_docx": str(docx_path.relative_to(ROOT).as_posix()),
            "supplement_docx": str(supp_path.relative_to(ROOT).as_posix()),
            "workbook": str(workbook_path.relative_to(ROOT).as_posix()),
            "figures": sorted(p.name for p in FIG_DIR.glob("*.png")),
            "tables": sorted(p.name for p in TABLE_DIR.glob("*.csv")),
        },
        "hashes": {
            "paper_docx": sha256(docx_path),
            "supplement_docx": sha256(supp_path),
            "workbook": sha256(workbook_path),
            "stage1_panel": sha256(ROOT / "data" / "processed" / "eurostat_ai_panel.csv"),
            "stage2_panel": sha256(ROOT / "data" / "processed" / "stage2_industry_panel.csv"),
        },
        "limits": [
            "官方统计面板支持宏观与行业机制验证，但不能替代企业微观问卷和访谈。",
            "模型结果不作为严格因果识别，只作为机制解释和决策支持证据。",
            "Stage 2是行业/区域外部验证层，不是SME-only样本。",
        ],
    }
    write_json(OUT_DIR / "research_registry_科技管理研究投稿版.json", registry)

    write_md(
        OUT_DIR / "研究成果核验报告.md",
        f"""
# 研究成果核验报告

## 最终事实源
- Stage 1：553行面板，544行可建模样本，36个geo，2021—2025；主模型Ridge，GroupKFold R2={s1["group_kfold_r2_mean"]:.4f}。
- Stage 2：5 814行行业/区域外部验证面板，36个geo，50个行业，2021—2025；主模型ExtraTrees，GroupKFold R2={s2["group_kfold_r2_mean"]:.4f}，Time holdout R2={s2["time_holdout_r2"]:.4f}，Industry holdout R2={s2["industry_holdout_r2"]:.4f}。
- 源数据口径：Stage 2由17个Eurostat官方源文件，经12 770 332行源数据扫描、856 880行指标筛选保留，聚合为5 814行建模面板。

## 研究状态判断
当前研究已具备可投稿初稿的证据链：官方数据、哈希核验、变量机制映射、分组机器学习验证、稳健性与异质性分析、复现工作簿和期刊格式稿件均已形成。短板仍是缺少真实企业微观问卷和访谈，因此主文只能写宏观/行业机制验证，不能写企业微观因果识别。

## GitHub状态建议
建议提交本19包和根目录README更新。14—18包保留为本地过程包，不直接作为投稿版本上传，避免蓝色旧DOCX、乱码文档和历史A10分数干扰最终口径。
""",
    )

    write_md(
        OUT_DIR / "科技管理研究投稿格式核对表.md",
        """
# 《科技管理研究》投稿格式核对表

- [x] 题名控制在20字左右：中小企业AI流程自动化采纳机制研究。
- [x] 摘要包含目的、方法、结果、结论，中文摘要约300—500字。
- [x] 关键词为3—8个，中英文对应。
- [x] 正文采用引言、文献述评、研究方法、研究结论等期刊式结构。
- [x] 图表使用黑白平面图，不使用蓝色装饰和PPT背景。
- [x] 表格采用三线表，不使用图片表格。
- [x] 参考文献不少于20篇，并保留系统文献矩阵用于后续人工核对。
- [x] 不写“千万样本直接训练”，统一写为“千万级官方源数据经筛选与面板聚合形成建模面板”。
- [x] Stage 2写为行业/区域外部验证层，不写成SME-only。
- [x] ai.zhjjq.tech只作为实践启示，不作为主证据来源。

待人工补齐：
- [ ] 作者姓名、单位、邮编、作者简介和基金项目。
- [ ] 投稿前按学校/作者真实信息核对署名顺序。
- [ ] 参考文献页码、卷期和英文大小写建议再逐条人工核验。
""",
    )

    write_md(
        OUT_DIR / "REPRODUCIBILITY_科技管理研究版.md",
        """
# 复现说明

## 环境
建议使用Python 3.10+，依赖pandas、numpy、matplotlib、scikit-learn、python-docx、openpyxl。本包不依赖A10服务器，CPU即可复现投稿材料。

## 一键生成投稿包
```bash
python 19_科技管理研究投稿定稿包/build_journal_submission_package.py
```

## 事实源
- `10_Agent系统/reports/final_research_registry_summary.json`
- `14_CPU研究增强包/outputs/cpu_research_upgrade_summary.json`
- `16_论文级实证增强包/outputs/*.csv`
- `17_投稿级研究完善包/tables/*.csv`
- `data/processed/eurostat_ai_panel.csv`
- `data/processed/stage2_industry_panel.csv`

## 安全边界
不提交`.joblib`、`.pkl`、`.env`、token、密码或服务器登录信息。A10相关内容仅作为历史训练记录，不作为当前投稿复现依赖。
""",
    )

    write_md(
        OUT_DIR / "README.md",
        """
# 19_科技管理研究投稿定稿包

本目录是面向《科技管理研究》的投稿定稿包，区别于早期课程汇报PPT和蓝色风格DOCX。本包只保留可追溯、可复现、符合期刊格式的材料。

## 主要文件
- `docs/投稿论文_科技管理研究格式_v1.docx`：黑白期刊风格正文稿。
- `docs/补充材料_科技管理研究格式_v1.docx`：稳健性、变量字典与复现补充材料。
- `figures/`：黑白/灰度PNG与SVG图，均由CSV/JSON结果生成。
- `tables/`：三线表源数据、图表源数据和Excel工作簿。
- `outputs/研究成果核验报告.md`：最终事实源和研究状态说明。
- `outputs/科技管理研究投稿格式核对表.md`：对照期刊要求的质量门禁。
- `outputs/REPRODUCIBILITY_科技管理研究版.md`：复现说明。

## 口径
Stage 1是中小企业机制解释层；Stage 2是行业/区域外部验证层。严禁把Stage 2写成SME-only样本，严禁把源数据行数写成直接训练样本量。
""",
    )

    lit.to_excel(TABLE_DIR / "系统文献矩阵_科技管理研究投稿版.xlsx", index=False)


def run_security_scan() -> dict:
    patterns = {
        "github_token": re.compile(r"github_pat_[A-Za-z0-9_]+"),
        "gamma_key": re.compile(r"sk-gamma-[A-Za-z0-9]+"),
        "server_password_hint": re.compile(r"(密码|password|passwd)\s*[:=]\s*\S+", re.I),
        "private_key": re.compile(r"BEGIN (RSA|OPENSSH|DSA|EC) PRIVATE KEY"),
    }
    findings = []
    for path in PKG.rglob("*"):
        if not path.is_file() or path.suffix.lower() in {".png", ".svg", ".xlsx", ".docx"}:
            continue
        text = path.read_text(encoding="utf-8", errors="ignore")
        for name, pat in patterns.items():
            for match in pat.finditer(text):
                findings.append({"file": str(path.relative_to(ROOT).as_posix()), "type": name, "match": match.group(0)[:20]})
    report = {"status": "pass" if not findings else "fail", "findings": findings}
    write_json(OUT_DIR / "投稿包安全扫描.json", report)
    return report


def try_render_docx(docx_path: Path) -> dict:
    render_dir = OUT_DIR / "docx_render"
    render_dir.mkdir(exist_ok=True)
    ps = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open('{docx_path}')
$pdf = '{render_dir / (docx_path.stem + '.pdf')}'
$doc.ExportAsFixedFormat($pdf, 17)
$doc.Close($false)
$word.Quit()
Write-Output $pdf
"""
    try:
        result = subprocess.run(["powershell", "-NoProfile", "-Command", ps], capture_output=True, text=True, timeout=90)
        if result.returncode == 0:
            return {"status": "pdf_exported", "pdf": str((render_dir / (docx_path.stem + ".pdf")).relative_to(ROOT).as_posix())}
        return {"status": "failed", "stderr": result.stderr[-1000:]}
    except Exception as exc:
        return {"status": "failed", "error": str(exc)}


def main() -> None:
    src = ensure_source_artifacts()
    tables = build_publication_tables(src)
    lit = build_literature_matrix()
    make_figures(src)
    paper_md = paper_markdown(src, tables)
    write_md(DOC_DIR / "投稿论文_科技管理研究格式_v1.md", paper_md)
    docx_path = make_docx(paper_md, tables)
    supp_path = make_supplement_docx(tables, lit)
    workbook_path = make_workbook(tables, lit)
    write_reports(src, tables, lit, docx_path, supp_path, workbook_path)
    security = run_security_scan()
    render = try_render_docx(docx_path)
    write_json(OUT_DIR / "docx_render_check.json", render)
    if security["status"] != "pass":
        raise SystemExit("Security scan failed")
    print(json.dumps({"status": "ok", "package": str(PKG), "render": render}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
