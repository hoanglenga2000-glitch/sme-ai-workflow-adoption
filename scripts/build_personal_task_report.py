from __future__ import annotations

import re
import shutil
import subprocess
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
FINAL = ROOT / "课程最终提交材料"
OUT_DIR = FINAL / "04_个人作业总结" / "04_个人任务报告"
DOCX = OUT_DIR / "信管2301景浩伟202321054012个人任务报告.docx"
PDF = OUT_DIR / "信管2301景浩伟202321054012个人任务报告.pdf"
TABLES = ROOT / "outputs" / "tables"
REPORTS = ROOT / "outputs" / "reports"
FIGS = ROOT / "outputs" / "figures" / "academic"
BLACK = RGBColor(0, 0, 0)


def set_font(run, size=10.5, bold=False, font="微软雅黑"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = BLACK


def add_title(doc: Document, text: str, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=size, bold=True, font="微软雅黑")


def heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 3)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 14, bold=True, font="微软雅黑")


def para(doc: Document, text: str, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_font(r, size=10.5)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    set_font(r, size=9, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def table(doc: Document, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True)
        shade(t.rows[0].cells[i], "D9D9D9")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
    doc.add_paragraph()


def figure(doc: Document, filename: str, caption: str, width_cm=13.6):
    path = FIGS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_font(r, size=9)


def normalize_docx(path: Path):
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                text = data.decode("utf-8", "ignore")
                text = re.sub(r'(<w:color\b[^>]*\bw:val=")[^"]+(")', r"\g<1>000000\2", text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def fmt(x, digits=4):
    try:
        return f"{float(x):.{digits}f}"
    except Exception:
        return str(x)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for old in OUT_DIR.glob("信管2301景浩伟202321054012个人案例总结.*"):
        old.unlink()

    alg = pd.read_csv(TABLES / "course_algorithm_comparison.csv")
    reg = pd.read_csv(TABLES / "course_regression_summary.csv")
    quality = pd.read_csv(TABLES / "enhanced_data_quality_audit.csv")
    vif = pd.read_csv(TABLES / "course_vif_diagnostics.csv")
    stage2_imp = pd.read_csv(TABLES / "stage2_feature_importance.csv")
    sme_cluster = pd.read_csv(TABLES / "sme_persona_clusters.csv")
    stage2_cluster = pd.read_csv(TABLES / "stage2_persona_clusters.csv")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    styles = doc.styles
    for name in ["Normal", "Body Text"]:
        if name in styles:
            styles[name].font.name = "微软雅黑"
            styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            styles[name].font.color.rgb = BLACK

    add_title(doc, "机器学习课程个人任务报告")
    add_title(doc, "企业AI部署偏好与治理机制研究", size=18)
    table(
        doc,
        ["项目", "内容"],
        [
            ["姓名", "景浩伟"],
            ["班级", "信管2301"],
            ["学号", "202321054012"],
            ["小组角色", "负责人、组长、答辩人"],
            ["任务定位", "选题设计、研究框架、数据处理思路、特征工程、模型评估、报告整合、PPT逻辑和课堂答辩"],
            ["提交文件", "个人任务报告DOCX/PDF、十次个人作业PDF、平时作业汇总、小组最终展示材料"],
        ],
    )
    para(doc, "本报告根据本人在课程案例中的实际分工撰写，重点说明我在选题设计、数据处理、特征工程、机器学习建模、结果解释、GitHub整理和小组展示中的具体工作。报告内容与投稿稿、GitHub公开项目、最终课程报告和小组PPT保持一致，不把模型相关性写成因果结论，也不把私密问卷或访谈材料混入公开数据证据。")

    heading(doc, "1 任务背景与个人职责")
    para(doc, "本小组课程案例以“企业AI部署偏好与治理机制”为主题，研究中小企业在AI流程自动化采纳过程中的效率需求、安全顾虑和部署路径。我的工作不是单独完成某一页PPT，而是把案例从选题、数据、模型、结果解释到最终展示串成一条完整证据链。")
    para(doc, "作为负责人和答辩人，我承担整体选题设计、研究框架搭建、任务统筹与进度安排；同时承担主要技术工作，重点负责数据处理思路设计、特征工程、机器学习方法与模型评估部分，包括OLS、Ridge、随机森林、ExtraTrees、MLP等模型的选择与对比，GroupKFold交叉验证方法说明，R²、MAE等指标解释，以及最终报告整合、PPT逻辑梳理和课堂答辩展示。")
    table(
        doc,
        ["成员", "分工内容"],
        [
            ["景浩伟", "负责整体选题设计、研究框架、任务统筹、数据处理思路、特征工程、模型评估、结果解释、最终报告整合、PPT逻辑和答辩展示。"],
            ["张新通", "负责数据来源与数据生命周期部分，包括Eurostat官方数据说明、真实性校验、源数据下载记录、清洗流程整理以及Stage 1、Stage 2数据边界说明。"],
            ["刘子涵", "负责机器学习方法与模型评估部分，参与OLS、Ridge、随机森林、ExtraTrees、MLP等模型原理整理、指标对比和图表解释。"],
            ["黄陈熙", "负责机制解释、企业部署策略和PPT视觉整理部分，包括特征重要性解释、效率需求与安全顾虑分析、部署偏好策略矩阵和答辩辅助材料检查。"],
        ],
    )

    heading(doc, "2 项目工作流与GitHub证据链")
    para(doc, "本项目的工作流分为七步：确定研究问题、下载并登记官方数据、清洗与构建建模面板、设计变量体系、训练与验证模型、解释结果并形成部署建议、整理GitHub公开展示材料。最终仓库保留了数据、源码、PPT和报告，删除了历史草稿、临时预览图、缓存文件和不适合公开的私密材料。")
    table(
        doc,
        ["模块", "本地证据", "作用"],
        [
            ["数据", "课程最终提交材料\\01_数据", "保存processed、samples、raw manifests、结果表、报告和图表。"],
            ["源码", "课程最终提交材料\\02_源码", "保存数据下载、清洗、建模、图表生成、Agent原型和复现说明。"],
            ["汇报", "课程最终提交材料\\03_小组汇报PPT和报告", "保存18页小组最终汇报PPT、PDF、证据映射和最终课程报告。"],
            ["个人", "课程最终提交材料\\04_个人作业总结", "保存平时作业汇总、十次个人作业PDF和本个人任务报告。"],
        ],
    )
    para(doc, f"截至本次整理，仓库可见文件约321个，最终提交材料约145个文件。最终包不是原始工作区的简单复制，而是围绕课程验收顺序重新整理：先给数据，再给源码，再给小组展示材料，最后给个人作业与个人任务报告。")

    heading(doc, "3 数据来源、清洗与建模面板")
    para(doc, "本项目使用Eurostat官方企业ICT调查数据，所有源文件通过manifest记录来源、下载状态和校验信息。投稿稿中明确写明，Stage 1使用10个Eurostat数据集，原始数据134,367行，经长面板转换、目标变量非空过滤和覆盖率阈值筛选后，形成553个唯一企业规模组观测值，其中544个进入建模。Stage 2使用17个Eurostat数据集，官方源文件共12,770,332行，经源文件剖析、特征过滤和面板聚合后，形成5,814行建模面板。")
    para(doc, "我在报告和答辩中专门强调一个边界：项目不是直接把千万级原始行丢给模型训练，而是通过清洗、筛选和聚合得到可解释的面板数据。这个过程体现了数据挖掘课程中比“跑模型”更基础的部分，即数据来源核验、字段口径统一、缺失处理、目标变量过滤和特征泄漏控制。")
    table(
        doc,
        ["数据层", "建模行数", "列数", "目标非空", "国家/地区", "年份范围", "使用说明"],
        [
            [
                "Stage 1：SME规模层" if "stage1" in row["dataset"] else "Stage 2：行业/区域GE10",
                int(row["rows"]),
                int(row["columns"]),
                int(row["target_nonnull"]),
                int(row["geo_count"]),
                f"{int(row['year_min'])}-{int(row['year_max'])}",
                "SME机制解释" if "stage1" in row["dataset"] else "外部验证，不写成SME-only",
            ]
            for _, row in quality.iterrows()
        ],
    )
    figure(doc, "fig1_academic_validation_clean.png", "图1 数据验证与建模流程")

    heading(doc, "4 特征工程与变量体系")
    para(doc, "变量体系来自投稿论文中的TOE框架和TAM模型。技术维度对应AI能力、云服务能力、数字基础和数据成熟度；组织维度对应ICT人才、培训、治理成熟度和部署准备度；环境维度对应国家、行业、年份和市场数字化差异。TAM中的感知有用性被转写为效率需求，感知易用性被转写为部署准备度。")
    para(doc, "在具体建模时，因变量为E_AI_TPA，即企业使用AI进行工作流自动化的比例。核心自变量包括机器学习能力、自然语言生成、部署准备度、安全顾虑、治理成熟度、数字基础、数据成熟度等。为了保证模型结果可信，我在整理报告时明确排除了目标变量及其直接派生字段，避免模型使用同义指标获得虚高分数。")
    table(
        doc,
        ["变量类型", "代表指标", "任务含义"],
        [
            ["因变量", "E_AI_TPA", "企业使用AI进行工作流自动化处理的比例。"],
            ["效率需求", "E_AI_TML、E_AI_TNLG、E_AI_CC", "体现企业是否具备AI技术使用基础以及对效率提升的需求。"],
            ["安全顾虑", "security_concern_index", "衡量隐私、安全、合规和数据管理约束。"],
            ["部署准备度", "deployment_readiness_index、E_CC_PDEV、E_CC_DA", "衡量云服务、数据分析和流程接入能力。"],
            ["治理成熟度", "governance_maturity_proxy", "体现培训、数据管理和流程治理基础。"],
            ["控制变量", "geo、year、nace_r2、size_emp", "控制国家、年份、行业和规模差异。"],
        ],
    )

    heading(doc, "5 机器学习方法与模型评估")
    para(doc, "我负责把课程中的算法方法放入同一个研究流程。OLS用于解释变量方向和显著性，Ridge用于缓解共线性，随机森林和ExtraTrees用于捕捉非线性关系，MLP作为神经网络基线，KMeans用于企业画像聚类。模型评价主要使用R²和MAE，并采用按国家分组的GroupKFold交叉验证。")
    para(doc, "GroupKFold是本项目最关键的验证设计。企业ICT数据具有明显的国家结构，如果随机划分训练集和测试集，同一国家的相似观测可能同时出现在两边，导致分数偏高。按国家分组可以降低地理信息泄漏，使模型评价更接近外部验证。")
    table(
        doc,
        ["数据层", "模型", "R²均值", "MAE均值", "特征数", "验证方式"],
        [[r["dataset"], r["model_cn"], fmt(r["r2_mean"]), fmt(r["mae_mean"]), int(r["feature_count"]), r["validation"]] for _, r in alg.iterrows()],
    )
    table(
        doc,
        ["数据层", "n", "R²", "Adj.R²", "特征数", "p<0.05变量数"],
        [[r["dataset"], int(r["n"]), fmt(r["r2"]), fmt(r["adj_r2"]), int(r["features"]), int(r["significant_05"])] for _, r in reg.iterrows()],
    )
    figure(doc, "fig1a_model_comparison_ppt.png", "图2 模型交叉验证比较")

    heading(doc, "6 结果解释与论文衔接")
    para(doc, "从模型结果看，Stage 1中Ridge模型在GroupKFold by country下R²均值为0.8744，MAE均值为1.7730；Stage 2中ExtraTrees在行业/区域外部验证口径下R²均值为0.7073，MAE均值为2.1060。两个阶段的指标口径不同，但共同说明机器学习能力、部署准备度和治理成熟度是解释AI流程自动化采纳的重要线索。")
    para(doc, "投稿论文中还报告了完整训练集口径下Ridge R²=0.8680、MAE=1.8342，以及Stage 2锁定基准ExtraTrees GroupKFold R²=0.7245、MAE=1.9646。课程报告采用这些结果时，我专门保留了口径说明，避免把不同验证方式混在一起，也避免把行业/区域GE10面板写成SME-only结论。")
    table(
        doc,
        ["诊断项", "结果或处理方式", "我的解释"],
        [
            ["VIF共线性", f"最高VIF示例：{vif.sort_values('vif', ascending=False).iloc[0]['feature_label']}={fmt(vif.sort_values('vif', ascending=False).iloc[0]['vif'], 2)}", "单个OLS系数不能孤立解释，需要结合Ridge和树模型结果。"],
            ["特征重要性", f"Stage 2最高特征：{stage2_imp.iloc[0]['feature']}，重要性{fmt(stage2_imp.iloc[0]['importance_mean'], 3)}", "机器学习能力是跨阶段稳定核心变量。"],
            ["聚类画像", f"Stage 1形成4类画像，Stage 2形成{stage2_cluster.shape[0]}类行业/区域画像", "聚类结果用于部署策略建议，不写成因果分群。"],
            ["证据边界", "访谈和问卷用于构念提出，不公开逐行原始数据", "公开证据以Eurostat、结果表、图表和代码为主。"],
        ],
    )
    figure(doc, "fig2_stage2_external_importance.png", "图3 Stage 2外部验证特征重要性")

    heading(doc, "7 企业部署策略与课堂展示")
    para(doc, "模型结果最终要服务展示和管理解释。我的答辩思路是先讲清楚数据链，再讲模型验证，最后把结果转成企业部署建议。安全顾虑不一定阻止AI采纳，而会改变部署路径：安全敏感型企业更适合本地化、私有云或混合部署；效率需求强但治理基础一般的企业，可以先从标准SaaS或API接入开始；治理成熟企业可以推进流程级自动化和跨系统集成。")
    para(doc, "小组最终PPT已经按18页展示逻辑整理：封面、研究目标、理论框架、数据生命周期、Stage 1和Stage 2模型结果、特征权重、部署演化路径、Agent落地、研究价值和方法边界。PPT证据映射表逐页对应数据表、报告或图表，避免展示材料和仓库证据脱节。")
    table(
        doc,
        ["最终展示材料", "数量/页数", "质量控制"],
        [
            ["小组最终汇报PPT", "18页", "PPTX可编辑，PDF可预览，已建立逐页证据映射。"],
            ["最终课程报告", "18页", "DOCX/PDF同步，字体黑色，无修订痕迹。"],
            ["个人十次作业", "10份PDF", "按01到10排序，保留代码、模型结果和图表页。"],
            ["质量核验清单", "1份MD", "记录页数、数据目录非空、图表分辨率、模型边界和测试结果。"],
        ],
    )

    heading(doc, "8 个人收获、问题与改进")
    para(doc, "这次任务让我意识到，机器学习课程项目不能只停留在模型名称和结果分数上。一个完整项目至少要回答五个问题：数据从哪里来，清洗规则是什么，特征为什么这样构建，验证方式是否会造成泄漏，结论能不能被复核。以前做单次作业时，我更多关注代码是否跑通；这次结课案例要求我把代码、数据、报告、PPT和GitHub放在同一条证据链里。")
    para(doc, "项目中也暴露出不足。第一，Eurostat指标口径复杂，变量解释需要持续对照官方说明；第二，VIF较高的变量需要谨慎解释，不能只看显著性；第三，PPT表达容易为了展示效果而写得过满，因此后期需要把“绝对”“完美”等词改成更稳的研究表述；第四，公开仓库和本地最终材料之间要区分清楚，未提交到GitHub的本地整理不能直接说成远端已经更新。")
    para(doc, "后续如果继续完善，我会优先做三件事：补充中国本土企业数据或公开统计数据，进一步验证欧洲样本结论能否迁移；增加时间留出或行业留出验证，检查模型在更严格场景下的表现；把Agent原型与真实业务流程连接，验证证据约束型回答能否帮助企业做部署选择。")

    heading(doc, "9 总结")
    para(doc, "本次个人任务围绕“企业AI部署偏好与治理机制研究”展开，我的核心贡献是把选题、数据、模型、报告和展示整合为可复核的课程项目。报告中的数据统计、模型指标和结论边界均来自GitHub项目输出和投稿论文基础材料。最终形成的个人任务报告不再只是学习感想，而是对本人实际承担工作的完整记录：有数据链、有代码链、有模型链、有展示链，也有对方法局限的说明。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("信管2301 景浩伟 202321054012｜机器学习课程个人任务报告")
    set_font(run, size=9)

    doc.save(DOCX)
    normalize_docx(DOCX)

    if PDF.exists():
        PDF.unlink()
    soffice = shutil.which("soffice") or r"C:\Users\景浩伟\.codex-office\bin\soffice.exe"
    subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", str(OUT_DIR), str(DOCX)], check=True)


if __name__ == "__main__":
    build()
    print(DOCX)
