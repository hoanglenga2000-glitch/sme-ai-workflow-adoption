from __future__ import annotations

import csv
import json
import re
import shutil
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
FINAL = REPO / "课程最终提交材料"
GROUP = FINAL / "03_小组汇报PPT和报告"
TABLES = REPO / "outputs" / "tables"
REPORTS = REPO / "outputs" / "reports"
FIGS = REPO / "outputs" / "figures" / "academic"
DOCX = GROUP / "中小企业AI流程自动化采纳机制研究案例_小组课程汇报报告.docx"
PDF = GROUP / "中小企业AI流程自动化采纳机制研究案例_小组课程汇报报告.pdf"
BLACK = RGBColor(0, 0, 0)


def read_csv(name: str) -> list[dict[str, str]]:
    with (TABLES / name).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


ALG = read_csv("course_algorithm_comparison.csv")
REG = read_csv("course_regression_summary.csv")
VIF = read_csv("course_vif_diagnostics.csv")
CLEAN = {row["stage"]: row for row in read_csv("cleaning_retention_summary.csv")}
STAGE2 = read_json(REPORTS / "stage2_large_model_metrics.json")


def set_run(run, *, size=10.5, bold=False, font="宋体") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = BLACK


def para(doc: Document, text: str = "", *, indent=True, size=10.5, bold=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold)


def heading(doc: Document, text: str, level=1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, size=15 if level == 1 else 13, bold=True, font="黑体")


def cell_text(cell, text: str, *, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    set_run(r, size=9.3, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell_text(t.rows[0].cells[i], header, bold=True)
        shade(t.rows[0].cells[i], "D9D9D9")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            cell_text(cells[i], value, align=align)
    doc.add_paragraph()


def figure(doc: Document, filename: str, caption: str, width=12.8) -> None:
    path = FIGS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run(r, size=9.2)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)
    for style_name in ["Normal", "Body Text"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "宋体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.color.rgb = BLACK


def cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("中小企业AI流程自动化采纳机制研究案例")
    set_run(r, size=20, bold=True, font="黑体")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("小组研究工作总结报告")
    set_run(r, size=16, bold=True, font="黑体")
    doc.add_paragraph()
    table(
        doc,
        ["项目", "内容"],
        [
            ["课程", "机器学习"],
            ["班级", "信管2301"],
            ["成员", "景浩伟、张新通、黄陈熙、刘子涵"],
            ["主题", "围绕中小企业AI流程自动化采纳，解释效率需求、安全顾虑、部署准备度与治理基础的关系"],
        ],
    )
    para(
        doc,
        "说明：本报告写的是小组研究工作的过程和结果。前期投稿稿只提供理论框架和问题意识参考，报告正文不按投稿论文写，也不展开最终提交材料清单。",
        indent=False,
    )
    doc.add_page_break()


def first(dataset: str, model: str) -> dict[str, str]:
    for row in ALG:
        if row["dataset"].startswith(dataset) and row["model_cn"] == model:
            return row
    raise KeyError((dataset, model))


def backup_current() -> None:
    backup = REPO / f"小组报告工作总结重写备份_{datetime.now():%Y%m%d_%H%M%S}"
    backup.mkdir(exist_ok=True)
    for path in [DOCX, PDF]:
        if path.exists():
            shutil.copy2(path, backup / path.name)


def normalize_docx_colors(path: Path) -> None:
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                text = data.decode("utf-8", "ignore")
                text = re.sub(r'(<w:color\b[^>]*\bw:val=")[^"]+(")', r"\g<1>000000\2", text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def build_report() -> None:
    ridge = first("Stage 1", "岭回归")
    et = first("Stage 2", "ExtraTrees")
    doc = Document()
    configure(doc)
    cover(doc)

    heading(doc, "一、我们为什么选择这个题目")
    para(doc, "小组最开始讨论选题时，没有想把AI写成一个很大的概念。课程项目需要能落到数据、模型和结果解释上，所以我们把问题收窄到“中小企业AI流程自动化采纳”。这个问题比较具体：企业不是简单地决定用不用AI，还要决定先用在哪些流程、采用SaaS还是API接入、是否需要本地化或混合部署。")
    para(doc, "这个题目和机器学习课程也能对应起来。我们可以把AI流程自动化使用率作为目标变量，把企业的AI能力、云服务基础、数据管理、安全顾虑、治理成熟度等因素作为解释变量，再用回归、树模型和聚类去检验哪些因素更稳定。这样，报告不是只讲观点，而是把观点放进一条可复核的数据链里。")
    para(doc, "前期投稿稿帮助我们确定了TOE框架、TAM思路以及“效率需求、安全顾虑、部署偏好”这几个概念。但课程报告不照搬投稿稿。我们把它改成小组研究工作总结，重点交代我们怎么做、做到了什么、哪些结论能说、哪些结论不能说。")
    doc.add_page_break()

    heading(doc, "二、研究问题怎样被拆成可建模任务")
    para(doc, "我们把研究问题拆成三层。第一层是效率需求：企业是否因为流程耗时、重复录入、人工复核等问题，希望用AI提高处理效率。第二层是安全顾虑：数据隐私、合规要求和模型可解释性会不会改变企业的选择。第三层是部署偏好：企业最终选择轻量SaaS、API接入、本地化部署，还是混合架构。")
    para(doc, "这三层不是彼此独立的。效率需求强，并不代表企业一定会马上深度采纳AI；安全要求高，也不代表企业完全拒绝AI。更合理的解释是：效率需求推动企业进入试点，安全和治理基础决定项目能不能继续往关键流程走，部署准备度影响AI最终能落到多深。")
    table(
        doc,
        ["研究概念", "在项目中的处理方式", "为什么这样处理"],
        [
            ["效率需求", "用AI能力、自然语言生成、认知计算等指标间接反映", "这些指标能说明企业是否已经具备把AI用于流程处理的基础"],
            ["安全顾虑", "用隐私、安全、数据管理和合规相关指标构造解释变量", "安全不是简单阻力，更像部署路径的选择条件"],
            ["部署准备度", "用云开发、云数据分析和数字基础指标表示", "企业是否能把AI接入业务系统，取决于已有数字基础"],
            ["治理成熟度", "用培训、流程开发、数据治理实践等指标代理", "AI进入关键流程后，需要组织规则和责任边界支撑"],
        ],
    )
    doc.add_page_break()

    heading(doc, "三、数据工作：最费时间的部分不是跑模型")
    para(doc, "真正花时间的是数据。Eurostat企业ICT数据的好处是来源公开、口径清楚、可以复现；麻烦也很明显，不同数据集的年份、国家、行业、企业规模和指标编码并不完全一致。我们不能把这些数据直接丢进模型，必须先把字段口径对齐。")
    para(doc, f"Stage 1使用企业规模组口径，官方多源数据共有{int(CLEAN['stage1_official_multisource']['raw_or_long_rows']):,}行，最后进入模型的是{int(CLEAN['stage1_official_multisource']['model_rows']):,}行。Stage 2用于外部验证，官方源数据链记录{int(CLEAN['stage2_large_sources_profiled']['raw_or_long_rows']):,}行，经过指标筛选和面板聚合后形成{int(CLEAN['stage2_large_sources_profiled']['model_rows']):,}行建模数据。这里要说清楚：我们不是直接用千万级原始行训练模型，而是先做清洗、筛选和聚合。")
    para(doc, "清洗过程主要处理四件事：统一国家、年份、行业和规模字段；把长表转成可建模的宽面板；剔除覆盖率太低或含义不稳定的指标；去掉与目标变量过近的字段，避免信息泄漏。这个过程看起来不如模型结果醒目，但它决定了后面的R²、MAE和特征重要性有没有意义。")
    table(
        doc,
        ["阶段", "原始/扫描记录", "中间面板", "建模样本", "用途"],
        [
            ["Stage 1", f"{int(CLEAN['stage1_official_multisource']['raw_or_long_rows']):,}", f"{int(CLEAN['stage1_official_multisource']['panel_rows']):,}", f"{int(CLEAN['stage1_official_multisource']['model_rows']):,}", "SME规模组机制解释"],
            ["Stage 2", f"{int(CLEAN['stage2_large_sources_profiled']['raw_or_long_rows']):,}", f"{int(CLEAN['stage2_large_sources_profiled']['panel_rows']):,}", f"{int(CLEAN['stage2_large_sources_profiled']['model_rows']):,}", "行业/区域外部验证"],
            ["Stage 2指标过滤", f"{int(CLEAN['stage2_indicator_filtering']['raw_or_long_rows']):,}", f"{int(CLEAN['stage2_indicator_filtering']['panel_rows']):,}", f"{int(CLEAN['stage2_indicator_filtering']['model_rows']):,}", "从候选指标压缩到可解释特征"],
        ],
    )
    figure(doc, "fig1_academic_validation_clean.png", "图1 数据清洗与建模流程")
    doc.add_page_break()

    heading(doc, "四、变量体系：把管理问题变成模型特征")
    para(doc, "变量设计时，我们没有只追求数量。企业ICT数据里有很多看起来相关的指标，如果全部放进模型，结果反而不好解释。小组最后保留的是能对应研究问题的变量组：AI能力、云服务能力、数据成熟度、安全顾虑、部署准备度、治理成熟度，以及国家、年份、行业和规模控制变量。")
    para(doc, "因变量是AI流程自动化使用率。这个变量能直接对应课程汇报的主题，也能把“AI采纳”从泛泛而谈变成一个可度量目标。为了不让模型偷看答案，我们在训练前剔除了目标变量及其直接派生字段。这个处理让分数更可信，也更容易向老师解释。")
    table(
        doc,
        ["变量组", "代表指标", "解释"],
        [
            ["因变量", "E_AI_TPA", "企业使用AI进行工作流自动化处理的比例"],
            ["AI能力", "E_AI_TML、E_AI_TNLG、E_AI_CC", "企业是否已经使用机器学习、自然语言生成、认知计算等AI能力"],
            ["部署准备度", "E_CC_PDEV、E_CC_DA", "企业是否具备云开发、云数据分析和系统接入基础"],
            ["治理成熟度", "governance_maturity_proxy", "培训、流程开发、数据管理等组织基础"],
            ["安全顾虑", "security_concern_index", "隐私、安全、合规和数据管理约束"],
            ["控制变量", "geo、year、nace_r2、size_emp", "国家、年份、行业和规模差异"],
        ],
    )
    doc.add_page_break()

    heading(doc, "五、模型设计：为什么不用一个模型讲到底")
    para(doc, "我们没有只选一个模型。OLS便于看变量方向和显著性，Ridge能处理共线性，Random Forest和ExtraTrees能捕捉非线性关系，KMeans能把结果转成企业画像。每个模型的任务不一样，放在一起才能支撑一个比较完整的课程案例。")
    para(doc, "验证方法上，我们使用GroupKFold by country。这样做是因为企业ICT数据天然带有国家结构，同一国家的观测很可能共享制度环境、数字基础和产业特点。如果随机划分训练集和验证集，模型可能只是记住国家差异，分数会偏高。按国家分组验证更严格，也更能说明模型是否有外部解释力。")
    table(
        doc,
        ["方法", "我们用它回答的问题", "课程知识点"],
        [
            ["OLS", "变量方向是否符合研究预期", "多元线性回归、显著性解释"],
            ["Ridge", "在共线性存在时结果是否更稳", "正则化回归"],
            ["Random Forest / ExtraTrees", "变量之间是否存在非线性关系", "集成学习、特征重要性"],
            ["GroupKFold", "模型是否只是记住国家差异", "交叉验证、数据泄漏控制"],
            ["KMeans", "不同企业画像怎样对应部署路径", "无监督学习、聚类解释"],
        ],
    )
    figure(doc, "fig1a_model_comparison_ppt.png", "图2 模型交叉验证结果")
    doc.add_page_break()

    heading(doc, "六、主要结果：哪些结论能站得住")
    para(doc, f"Stage 1聚焦中小企业规模组。Ridge模型在GroupKFold by country下取得R²均值{float(ridge['r2_mean']):.4f}、MAE均值{float(ridge['mae_mean']):.4f}。这个结果说明，在SME规模组口径下，AI能力、部署准备度和治理相关变量对AI流程自动化采纳差异有较强解释力。")
    para(doc, f"Stage 2转向行业和区域口径。ExtraTrees在GroupKFold by country下取得R²均值{float(et['r2_mean']):.4f}、MAE均值{float(et['mae_mean']):.4f}。这个分数比Stage 1低，但场景也更复杂：它覆盖36个国家或地区和50个NACE行业，不再是单纯的SME规模组解释。我们把它写成外部验证，而不是写成SME-only结论。")
    para(doc, "两个阶段共同指向一个判断：机器学习能力是比较稳定的核心变量，部署准备度和治理成熟度会影响AI能否真正进入流程。安全顾虑没有被我们写成“越高越不采纳”，因为数据和业务逻辑都提示它更像路径选择条件。安全要求越高，企业越可能从轻量工具转向更可控的本地化、私有云或混合部署。")
    table(
        doc,
        ["阶段", "模型", "R²均值", "MAE均值", "验证方式"],
        [
            [row["dataset"], row["model_cn"], f"{float(row['r2_mean']):.4f}", f"{float(row['mae_mean']):.4f}", row["validation"]]
            for row in ALG
        ],
    )
    figure(doc, "fig1b_sme_importance_ppt.png", "图3 Stage 1特征重要性")
    figure(doc, "fig2_stage2_external_importance.png", "图4 Stage 2特征重要性")
    doc.add_page_break()

    heading(doc, "七、结果解释：从模型分数回到企业场景")
    para(doc, "如果只看模型分数，这个项目就只完成了一半。我们更关心的是这些结果怎样解释企业采纳路径。对中小企业来说，AI流程自动化往往不是一次性大规模替换，而是从低风险流程开始试点，再逐步进入客服、财务、供应链、文档处理等业务环节。")
    para(doc, "效率需求强的企业更容易启动项目，但如果数据分散、接口混乱、治理责任不清，AI很可能停在试点。安全敏感的企业也不是不能用AI，而是需要更清楚的权限、日志、数据边界和模型解释。部署偏好由此变成一个很实际的问题：什么时候用SaaS，什么时候接API，什么时候必须本地化，什么时候适合混合架构。")
    table(
        doc,
        ["企业画像", "典型特征", "更合适的部署路径"],
        [
            ["效率驱动型", "流程重复、人工复核多，安全要求相对可控", "先用SaaS或API接入，快速验证ROI"],
            ["治理成熟型", "数据管理和培训基础较好，流程标准化程度高", "推进流程级自动化和跨系统集成"],
            ["安全敏感型", "合规压力大，数据边界严格", "本地化、私有云或混合部署，配合审计机制"],
            ["基础薄弱型", "系统分散、数据质量不稳", "先补数据治理和流程标准化，再扩大AI应用"],
        ],
    )
    doc.add_page_break()

    heading(doc, "八、小组分工与协作过程")
    para(doc, "这次小组工作不是把任务简单拆成几页PPT。比较核心的工作是把研究问题、数据处理、模型结果和展示语言统一起来。每个人负责的部分不同，但最后必须能接成一条线：问题从哪里来，数据怎么证明，模型怎么验证，结果怎样解释。")
    table(
        doc,
        ["成员", "主要工作", "在汇报中的体现"],
        [
            ["景浩伟", "选题统筹、研究框架、数据与模型流程、报告整合、PPT逻辑和答辩", "把投稿思路改成课程案例，控制研究边界和整体叙事"],
            ["张新通", "数据来源、数据生命周期、源数据真实性和清洗流程整理", "说明为什么数据可信、为什么需要从原始记录聚合到建模面板"],
            ["刘子涵", "模型方法、指标解释、模型结果核对和图表解释", "解释Ridge、ExtraTrees、R²、MAE和GroupKFold结果"],
            ["黄陈熙", "机制解释、部署策略、PPT表达和展示节奏检查", "把模型结果转成企业画像、部署建议和课堂讲述语言"],
        ],
    )
    para(doc, "协作中最容易出问题的是口径。比如Stage 2不是SME-only，千万级源数据不能说成直接训练数据，模型相关性不能写成因果证明。后期我们反复改的重点，就是把这些边界写清楚。")
    doc.add_page_break()

    heading(doc, "九、我们遇到的问题和处理方式")
    table(
        doc,
        ["问题", "具体表现", "处理方式"],
        [
            ["数据口径复杂", "不同数据集年份、国家、行业和指标覆盖不一致", "先做源文件剖析，再统一字段并按面板聚合"],
            ["变量共线性", "云能力、数据成熟度、部署准备度之间相关性较高", "同时使用VIF、Ridge和树模型重要性，不孤立解释单个OLS系数"],
            ["展示容易写过满", "初稿里有偏论文式、偏宣传式的表述", "改成课程汇报语言，只保留能被数据和模型支撑的说法"],
            ["公开边界", "访谈、问卷和公开统计数据不能混成同一类证据", "公开证据以Eurostat、清洗结果、模型表和图表为主"],
        ],
    )
    para(doc, "这些问题也让我们意识到，机器学习项目不只是训练模型。数据边界、变量含义、验证方式和表达口径都会影响最终质量。一个分数很好看的模型，如果数据泄漏没有处理，或者结论写过头，反而会降低可信度。")
    doc.add_page_break()

    heading(doc, "十、课程收获与后续改进")
    para(doc, "从课程学习角度看，这个项目把几类算法放进了同一个真实问题中。OLS和Ridge让我们看到线性模型的解释价值，树模型让我们看到非线性关系，KMeans让我们把结果转成画像，GroupKFold提醒我们验证方式会影响模型分数。")
    para(doc, "后续如果继续做，我们会优先补三件事。第一，加入中国本土中小企业数据，检查欧洲公开数据得到的结论能否迁移。第二，做更严格的时间留出或行业留出验证，看看模型面对新年份、新行业时是否还稳定。第三，把Agent原型接入更具体的业务场景，用真实案例验证部署建议是否有帮助。")
    para(doc, "本次小组汇报最终要表达的是：我们不是为了证明AI一定能解决中小企业所有问题，而是用公开数据和机器学习方法，把中小企业采纳AI流程自动化时的效率、安全、部署和治理条件讲清楚。这个结论克制一点，但更像一次真正完成过的数据研究。")
    doc.add_page_break()

    heading(doc, "十一、案例汇报时我们准备讲清楚的主线")
    para(doc, "课堂汇报不是把报告内容从头念一遍。我们准备的讲述主线是“一个问题、两层数据、三类机制、四种画像”。一个问题，是中小企业为什么会在AI流程自动化上出现采纳差异。两层数据，是Stage 1的SME规模组面板和Stage 2的行业/区域外部验证。三类机制，是效率需求、安全顾虑和部署准备度。四种画像，是效率驱动型、治理成熟型、安全敏感型和基础薄弱型企业。")
    para(doc, "这样讲的好处是，听众不用先理解所有变量名，也能顺着业务问题听下去。先讲企业为什么需要AI流程自动化，再讲我们怎么从公开数据里找到可计算指标，然后讲模型结果是否支持这个解释，最后回到企业应该怎么选部署路径。PPT里的每一部分都围绕这条线展开，而不是为了展示模型而展示模型。")
    para(doc, "答辩时最需要注意的是边界。我们会明确说明，Stage 1是SME规模组机制解释，Stage 2是行业/区域口径的外部验证；机器学习模型解释的是相关性和预测能力，不是因果证明；Eurostat数据来自欧洲企业环境，不能直接替代中国本土企业数据。把这些话说在前面，反而能让报告更可信。")
    table(
        doc,
        ["汇报环节", "要讲清楚的问题", "对应研究工作"],
        [
            ["研究动机", "为什么关注AI流程自动化而不是泛泛讲AI", "把AI采纳收窄到可度量目标变量"],
            ["数据处理", "为什么千万级源数据最后变成几千行面板", "字段统一、覆盖率过滤、面板聚合和泄漏控制"],
            ["模型验证", "模型分数是否可信", "GroupKFold by country、R²、MAE、VIF和特征重要性"],
            ["管理解释", "企业应该如何理解结果", "效率、安全、部署准备度和治理成熟度共同决定落地路径"],
        ],
    )
    doc.add_page_break()

    heading(doc, "十二、个人作业和小组案例之间的关系")
    para(doc, "平时作业不是和小组案例割裂的。十次作业里做过的数据预处理、回归、分类、树模型、聚类和模型评价，最后都在这个案例里找到了位置。前面的单次作业像是练习单个工具，小组案例则要求我们把这些工具放进一个完整问题中。")
    para(doc, "个人总结部分也不是简单写感想，而是回到自己在小组工作中的具体贡献。比如，数据来源和生命周期整理对应数据处理能力；模型指标解释对应课程算法理解；PPT逻辑和答辩准备对应把模型语言转成管理语言的能力。老师看个人总结时，应该能看到每个成员确实参与了研究过程。")
    para(doc, "这也是我们后期整理材料时坚持保留每个成员平时作业汇总和个人案例总结的原因。平时作业证明课程训练过程，个人总结说明本人在小组案例中的实际工作，小组报告则把四个人的工作合成一条完整研究线。三者放在一起，才比较接近这门课的评分要求。")
    table(
        doc,
        ["课程训练内容", "小组案例中的对应位置", "形成的能力"],
        [
            ["数据清洗与缺失处理", "Eurostat多源数据整理、面板聚合", "把原始数据变成可建模数据"],
            ["回归与正则化", "OLS、Ridge和VIF诊断", "解释变量方向和共线性问题"],
            ["树模型与集成学习", "Random Forest、ExtraTrees", "处理非线性关系和特征重要性"],
            ["聚类分析", "企业画像与部署路径", "把模型结果转成管理解释"],
            ["模型评估", "GroupKFold、R²、MAE", "判断模型结果是否稳健"],
        ],
    )
    doc.add_page_break()

    heading(doc, "十三、最终结论")
    para(doc, "这次小组案例最后形成的判断可以概括为三句话。第一，中小企业采纳AI流程自动化不是简单的技术选择，而是效率压力、数字基础、安全要求和治理能力共同作用的结果。第二，公开统计数据虽然不能替代企业微观访谈，但可以帮助我们先建立一套可复核的量化参照。第三，机器学习模型的价值不只是给出分数，更重要的是帮助我们把复杂的企业采纳问题拆成可以解释、可以讨论、可以改进的部分。")
    para(doc, "如果把项目放回课程学习本身，我们最大的收获是把零散算法连成了完整流程。数据清洗决定了模型能不能用，验证方式决定了分数能不能信，变量解释决定了结果能不能讲清楚。小组汇报也因此不只是展示PPT，而是把数据、模型和管理解释放到同一个问题里。")
    para(doc, "后续继续完善时，我们不会优先增加更复杂的模型，而会先补充更贴近中国中小企业的本土数据，做更严格的时间和行业留出验证，并把Agent原型放进更具体的业务流程中测试。这样才能判断模型结论是否真的能帮助企业做AI流程自动化部署选择。")
    table(
        doc,
        ["结论维度", "本次研究得到的判断"],
        [
            ["数据层面", "官方公开数据可以支撑课程级复现，但需要清洗、聚合和口径说明"],
            ["模型层面", "Ridge和ExtraTrees给出了较稳定的解释结果，但不能写成因果证明"],
            ["管理层面", "效率需求推动试点，安全和治理条件影响部署深度"],
            ["课程层面", "平时算法训练最终落到一个完整数据研究案例中"],
        ],
    )
    doc.add_page_break()

    heading(doc, "附录：关键诊断结果")
    table(
        doc,
        ["数据层", "n", "R²", "Adj.R²", "特征数", "p<0.05变量数"],
        [[row["dataset"], row["n"], f"{float(row['r2']):.4f}", f"{float(row['adj_r2']):.4f}", row["features"], row["significant_05"]] for row in REG],
    )
    table(
        doc,
        ["数据层", "特征", "含义", "VIF"],
        [[row["dataset"], row["feature"], row["feature_label"], f"{float(row['vif']):.2f}"] for row in VIF[:8]],
    )
    para(doc, "附录只保留少量关键诊断结果，供老师核对模型解释边界。完整的运行结果和图表用于复核，不在正文中展开成文件展示。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("机器学习课程小组研究工作总结报告")
    set_run(r, size=9)
    doc.save(DOCX)
    normalize_docx_colors(DOCX)


def convert_pdf() -> None:
    if PDF.exists():
        PDF.unlink()
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(GROUP), str(DOCX)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    generated = GROUP / f"{DOCX.stem}.pdf"
    if generated != PDF:
        generated.replace(PDF)


def audit() -> dict:
    import fitz

    with fitz.open(PDF) as doc:
        blank_pages = []
        for idx, page in enumerate(doc, 1):
            text = page.get_text("text").strip()
            pix = page.get_pixmap(matrix=fitz.Matrix(0.12, 0.12), alpha=False)
            data = pix.samples
            nonwhite = 0
            for i in range(0, len(data), pix.n):
                if any(channel < 245 for channel in data[i : i + 3]):
                    nonwhite += 1
            ratio = nonwhite / max(1, pix.width * pix.height)
            if len(text) < 10 and ratio < 0.015:
                blank_pages.append(idx)
        pages = doc.page_count

    with zipfile.ZipFile(DOCX) as zf:
        xml = "\n".join(
            zf.read(name).decode("utf-8", "ignore")
            for name in zf.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
    colors = sorted(set(re.findall(r'<w:color\b[^>]*\bw:val="([^"]+)"', xml)))
    bad_colors = [c for c in colors if c.lower() not in {"000000", "auto"}]
    text = ""
    with zipfile.ZipFile(DOCX) as zf:
        text = zf.read("word/document.xml").decode("utf-8", "ignore")
    banned = [
        "最终提交目录",
        "01_数据",
        "02_源码",
        "03_小组汇报PPT和报告",
        "04_小组成员个人作业整理",
        "文件清单",
        "验收清单",
        "专家/评委",
        "完美印证",
    ]
    hits = [term for term in banned if term in text]
    result = {
        "ok": not blank_pages and not bad_colors and not hits and pages >= 10,
        "pages": pages,
        "blank_pages": blank_pages,
        "bad_colors": bad_colors,
        "banned_hits": hits,
    }
    (GROUP / "小组研究工作总结报告_质量核验.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return result


def main() -> None:
    backup_current()
    build_report()
    convert_pdf()
    result = audit()
    if not result["ok"]:
        raise RuntimeError(json.dumps(result, ensure_ascii=False, indent=2))
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
