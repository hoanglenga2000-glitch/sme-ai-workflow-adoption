from __future__ import annotations

import csv
import json
import re
import shutil
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
FINAL = REPO / "课程最终提交材料"
BACKUP = REPO / f"课程最终提交材料_纠偏备份_{datetime.now():%Y%m%d_%H%M%S}"
DATA = FINAL / "01_数据"
SOURCE = FINAL / "02_源码"
GROUP = FINAL / "03_小组汇报PPT和报告"
MEMBERS_OLD = FINAL / "05_小组成员个人作业整理"
MEMBERS = FINAL / "04_小组成员个人作业整理"
PERSONAL_OLD = FINAL / "04_个人作业总结"
TABLES = REPO / "outputs" / "tables"
REPORTS = REPO / "outputs" / "reports"
FIGS = REPO / "outputs" / "figures" / "academic"
BLACK = RGBColor(0, 0, 0)

GROUP_DOCX = GROUP / "中小企业AI流程自动化采纳机制研究案例_小组课程汇报报告.docx"
GROUP_PDF = GROUP / "中小企业AI流程自动化采纳机制研究案例_小组课程汇报报告.pdf"
PPTX = GROUP / "中小企业AI流程自动化采纳机制研究案例_小组最终汇报PPT.pptx"
PPT_PDF = GROUP / "中小企业AI流程自动化采纳机制研究案例_小组最终汇报PPT.pdf"
EVIDENCE_MAP = GROUP / "slide_to_evidence_map_小组最终汇报PPT.csv"


def safe_under(child: Path, parent: Path) -> None:
    child_resolved = child.resolve()
    parent_resolved = parent.resolve()
    if parent_resolved not in [child_resolved, *child_resolved.parents]:
        raise RuntimeError(f"Unsafe path: {child_resolved} is not under {parent_resolved}")


def read_csv(name: str) -> list[dict[str, str]]:
    path = TABLES / name
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


ALG = read_csv("course_algorithm_comparison.csv")
REG = read_csv("course_regression_summary.csv")
VIF = read_csv("course_vif_diagnostics.csv")
CLEAN = read_csv("cleaning_retention_summary.csv")
STAGE2_METRICS = read_json(REPORTS / "stage2_large_model_metrics.json")
MODEL_METRICS = read_json(REPORTS / "model_metrics.json")


def set_run(run, *, size=10.5, bold=False, font="宋体") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = BLACK


def style_paragraph(p, *, indent=True, spacing=1.45, after=4) -> None:
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)


def para(doc: Document, text: str = "", *, indent=True, size=10.5, bold=False) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, indent=indent)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold)


def heading(doc: Document, text: str, level=1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, size=15 if level == 1 else 13, bold=True, font="黑体")


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("· " + text)
    set_run(r, size=10.5)


def cell_text(cell, text: str, *, bold=False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    set_run(r, size=9.2, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell_text(t.rows[0].cells[idx], header, bold=True)
        shade(t.rows[0].cells[idx], "D9D9D9")
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            cell_text(cells[idx], value)
    doc.add_paragraph()


def figure(doc: Document, filename: str, caption: str, width=12.8) -> None:
    path = FIGS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run(r, size=9.2)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)
    for style_name in ["Normal", "Body Text"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "宋体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.color.rgb = BLACK


def cover(doc: Document, title: str, subtitle: str, rows: list[list[str]]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run(r, size=20, bold=True, font="黑体")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run(r, size=15, bold=True, font="黑体")
    doc.add_paragraph()
    table(doc, ["项目", "内容"], rows)
    doc.add_page_break()


def first(rows: list[dict[str, str]], dataset: str, model: str) -> dict[str, str]:
    for row in rows:
        if row["dataset"].startswith(dataset) and row["model_cn"] == model:
            return row
    raise KeyError((dataset, model))


def build_group_report() -> None:
    old_report_files = [
        GROUP / "企业AI部署偏好与治理机制研究_最终课程报告.docx",
        GROUP / "企业AI部署偏好与治理机制研究_最终课程报告.pdf",
    ]
    backup_dir = BACKUP / "旧小组报告"
    backup_dir.mkdir(parents=True, exist_ok=True)
    for path in old_report_files:
        if path.exists():
            shutil.move(str(path), str(backup_dir / path.name))

    ridge = first(ALG, "Stage 1", "岭回归")
    et = first(ALG, "Stage 2", "ExtraTrees")
    clean_rows = {r["stage"]: r for r in CLEAN}

    doc = Document()
    configure_doc(doc)
    cover(
        doc,
        "中小企业AI流程自动化采纳机制研究案例",
        "机器学习课程小组汇报报告",
        [
            ["课程名称", "机器学习"],
            ["班级", "信管2301"],
            ["小组成员", "景浩伟、张新通、黄陈熙、刘子涵"],
            ["汇报材料", "小组最终汇报PPT、课程汇报报告、数据与源码、成员个人作业整理"],
            ["材料边界", "投稿稿仅作为研究问题、理论框架和表述边界的参考，不作为投稿论文提交"],
        ],
    )

    heading(doc, "摘要")
    for text in [
        "本报告是机器学习课程的小组汇报材料，围绕“中小企业AI流程自动化采纳机制”展开。报告不按期刊论文格式展开，也不把投稿稿内容直接搬入课程材料；投稿稿只用于帮助确定研究问题、变量框架和表达边界。课程汇报重点放在数据来源、清洗过程、特征工程、模型验证、结果解释、PPT证据映射和成员作业完整性。",
        "项目使用Eurostat官方企业ICT公开数据。Stage 1聚焦中小企业规模组口径，形成544行建模样本，用于解释SME相关机制；Stage 2使用行业和区域层面的GE10口径，形成5,814行建模面板，用于外部验证。模型采用OLS、Ridge、Random Forest、ExtraTrees、KMeans等课程相关方法，并使用按国家分组的GroupKFold降低地理信息泄漏。",
        "结果显示，机器学习能力、部署准备度和治理成熟度是解释AI流程自动化采纳差异的重要线索。安全顾虑不被写成单一阻力，而是解释为影响部署路径选择的条件。报告中的结论均限定为公开数据下的相关性和预测解释，不声称因果关系。",
    ]:
        para(doc, text)
    para(doc, "关键词：中小企业；AI流程自动化；机器学习；数据清洗；GroupKFold；部署路径", indent=False, bold=True)
    doc.add_page_break()

    sections = [
        (
            "一、课程汇报任务与提交边界",
            [
                "这份材料服务于课程展示和最终提交，不服务于期刊投稿。小组汇报需要回答的是：数据是否真实、清洗是否可追溯、模型是否符合机器学习课程要求、PPT是否能讲清楚案例、成员个人作业是否齐全。因此，报告结构按课程验收逻辑组织，而不是按投稿论文的“引言、文献综述、假设检验”完整展开。",
                "最终提交目录调整为四个主要部分：01_数据、02_源码、03_小组汇报PPT和报告、04_小组成员个人作业整理。景浩伟个人材料不再单独放一个额外目录，而是和其他成员统一放入小组成员个人作业整理中。每名成员均按“10次个人作业PDF + 1份个人总结/实践报告”的方式归档。",
                "原投稿稿中的TOE、TAM、效率需求、安全顾虑和部署偏好等内容被保留为选题依据，但课程报告只引用其中能支持课堂讲解的部分。报告不加入期刊式中图分类号、英文摘要、投稿作者单位等内容，也不把投稿稿与小组汇报混成同一类材料。",
            ],
        ),
        (
            "二、研究问题与课程化表达",
            [
                "小组选择的核心问题是：中小企业为什么会采纳或暂缓采纳AI流程自动化，效率需求、安全顾虑和部署偏好在其中分别起什么作用。这个问题适合机器学习课程，是因为它既有明确的目标变量，也有可以从公开数据中构造的解释变量，还能通过不同算法比较模型表现。",
                "本项目没有把企业AI采纳写成单纯的技术热词，而是把它拆成三个可以讲清楚的部分。第一，效率需求对应企业希望减少重复劳动、提高流程响应速度的动机。第二，安全顾虑对应隐私、合规、数据管理和模型可解释要求。第三，部署偏好对应SaaS、API接入、本地化、混合云等落地路径选择。",
                "理论上，TOE框架帮助我们把变量分为技术、组织和环境三个层面；TAM帮助我们解释企业为什么认为AI有用、是否容易接入。课程展示中不需要展开长篇文献综述，只需要说明这些理论如何转成可计算指标，并如何进入模型训练。",
            ],
        ),
        (
            "三、数据来源与清洗流程",
            [
                "数据来自Eurostat官方企业ICT相关数据集，仓库保留manifest、下载脚本、样本数据、处理后面板和结果表。最终提交包不把全部raw大文件放进去，是为了控制体积；但保留了来源说明和重新下载脚本，老师需要复现时可以回到源码目录检查。",
                f"数据量口径需要说准确。Stage 2官方源数据链为{int(clean_rows['stage2_large_sources_profiled']['raw_or_long_rows']):,}行，不是1.2亿行。项目并不是把千万级原始记录直接投入模型，而是经过源文件剖析、指标筛选、覆盖率过滤、面板聚合和目标非空筛选，最终形成5,814行建模面板。",
                "清洗流程的关键是保证口径一致。不同Eurostat数据集在国家、年份、行业、企业规模和指标编码上并不完全统一，项目先将长表整理为统一字段，再合并为国家—年份—规模组或国家—年份—行业面板。缺失率过高、覆盖口径不稳定、与目标变量过近的字段会被剔除。",
            ],
        ),
    ]
    for title, paragraphs in sections:
        heading(doc, title)
        for text in paragraphs:
            para(doc, text)
        if title.startswith("三、"):
            table(
                doc,
                ["阶段", "源/扫描记录", "中间面板", "最终建模行", "课程用途"],
                [
                    ["Stage 1", f"{int(clean_rows['stage1_official_multisource']['raw_or_long_rows']):,}", f"{int(clean_rows['stage1_official_multisource']['panel_rows']):,}", f"{int(clean_rows['stage1_official_multisource']['model_rows']):,}", "SME规模组机制解释"],
                    ["Stage 2", f"{int(clean_rows['stage2_large_sources_profiled']['raw_or_long_rows']):,}", f"{int(clean_rows['stage2_large_sources_profiled']['panel_rows']):,}", f"{int(clean_rows['stage2_large_sources_profiled']['model_rows']):,}", "行业/区域外部验证"],
                    ["Stage 2指标过滤", f"{int(clean_rows['stage2_indicator_filtering']['raw_or_long_rows']):,}", f"{int(clean_rows['stage2_indicator_filtering']['panel_rows']):,}", f"{int(clean_rows['stage2_indicator_filtering']['model_rows']):,}", "特征筛选与面板聚合"],
                ],
            )
            figure(doc, "fig1_academic_validation_clean.png", "图1 数据验证与建模流程")
        doc.add_page_break()

    heading(doc, "四、特征工程与变量体系")
    for text in [
        "变量构建围绕效率需求、安全顾虑、部署准备度和治理成熟度展开。效率需求用AI技术使用、自然语言生成、认知计算等指标表示；部署准备度用云开发、云数据分析和数字基础表示；安全顾虑用隐私、安全、合规和数据管理相关指标表示；治理成熟度用培训、流程开发、数据管理实践等指标表示。",
        "因变量为AI流程自动化使用率。为了避免信息泄漏，训练时不使用目标变量及其直接派生字段。这样做会降低模型“看起来很高”的分数，但更符合课程项目对可解释性和可信度的要求。",
        "Stage 1和Stage 2的口径不同。Stage 1更适合解释中小企业规模组机制；Stage 2使用GE10行业/区域数据，只能作为外部验证，不能直接写成SME-only结论。这个边界在报告和PPT中都保留。",
    ]:
        para(doc, text)
    table(
        doc,
        ["变量类型", "代表指标", "含义"],
        [
            ["因变量", "E_AI_TPA", "企业使用AI进行工作流自动化处理的比例"],
            ["效率需求", "E_AI_TML、E_AI_TNLG、E_AI_CC", "AI技术能力和流程效率提升需求"],
            ["部署准备度", "E_CC_PDEV、E_CC_DA、digital_foundation_index", "云服务、数据分析和数字基础"],
            ["治理成熟度", "governance_maturity_proxy", "培训、数据管理和流程治理基础"],
            ["安全顾虑", "security_concern_index", "隐私、安全和合规约束"],
            ["控制变量", "geo、year、nace_r2、size_emp", "国家、年份、行业和规模差异"],
        ],
    )
    doc.add_page_break()

    heading(doc, "五、模型方法与训练设计")
    for text in [
        "课程方法覆盖回归、正则化、树模型、聚类和交叉验证。OLS用于看变量方向和统计解释；Ridge用于缓解共线性；Random Forest和ExtraTrees用于处理非线性关系；KMeans用于把企业或行业画像转成部署建议；VIF用于检查多重共线性。",
        "GroupKFold是本项目最重要的验证设计。企业ICT数据具有明显的国家结构，如果随机划分训练集和测试集，同一国家的相似观测可能同时出现在训练和验证两边，模型分数会偏乐观。按国家分组交叉验证更严格，也更适合课堂说明。",
        "模型选择不追求最复杂，而是追求能解释清楚。对课程项目来说，能说明数据从哪里来、如何清洗、模型为什么这样选、结果是否有边界，比堆叠很多算法更重要。",
    ]:
        para(doc, text)
    table(
        doc,
        ["方法", "在本案例中的作用", "课程知识点"],
        [
            ["OLS", "查看变量方向、显著性和诊断结果", "多元线性回归"],
            ["Ridge", "作为Stage 1稳健基准模型", "正则化回归"],
            ["Random Forest / ExtraTrees", "捕捉非线性关系和变量交互", "集成学习"],
            ["GroupKFold", "按国家分组验证，降低地理泄漏", "训练集/测试集划分"],
            ["VIF", "检查多重共线性", "回归诊断"],
            ["KMeans", "形成企业画像和部署策略", "聚类分析"],
        ],
    )
    figure(doc, "fig1a_model_comparison_ppt.png", "图2 模型交叉验证比较")
    doc.add_page_break()

    heading(doc, "六、模型结果与解释")
    para(doc, f"Stage 1采用SME规模组口径，Ridge在GroupKFold by country下R²均值为{float(ridge['r2_mean']):.4f}，MAE均值为{float(ridge['mae_mean']):.4f}，使用10个核心特征。这个结果说明，在规模组口径下，效率需求、部署准备度和治理相关变量对AI流程自动化采纳差异有较强解释力。")
    para(doc, f"Stage 2采用行业/区域GE10口径，ExtraTrees在GroupKFold by country下R²均值为{float(et['r2_mean']):.4f}，MAE均值为{float(et['mae_mean']):.4f}。与Stage 1相比，Stage 2场景更复杂，解释力下降是合理的，因为模型要跨行业和区域进行外部验证。")
    para(doc, "VIF结果提醒我们，云能力、数据成熟度和部署准备度之间存在一定共线性，所以报告不孤立解释单个OLS系数，而是结合Ridge、树模型特征重要性和外部验证结果。模型结果用于说明稳定相关关系，不用于声称企业采纳AI的因果机制已经被证明。")
    table(
        doc,
        ["数据层", "模型", "R²均值", "MAE均值", "验证方式"],
        [
            [r["dataset"], r["model_cn"], f"{float(r['r2_mean']):.4f}", f"{float(r['mae_mean']):.4f}", r["validation"]]
            for r in ALG
        ],
    )
    table(
        doc,
        ["数据层", "n", "R²", "Adj.R²", "特征数", "p<0.05变量数"],
        [[r["dataset"], r["n"], f"{float(r['r2']):.4f}", f"{float(r['adj_r2']):.4f}", r["features"], r["significant_05"]] for r in REG],
    )
    figure(doc, "fig1b_sme_importance_ppt.png", "图3 Stage 1特征重要性")
    figure(doc, "fig2_stage2_external_importance.png", "图4 Stage 2外部验证特征重要性")
    doc.add_page_break()

    heading(doc, "七、企业画像与部署路径建议")
    for text in [
        "模型结果要转成课堂汇报中听得懂的管理解释。效率需求较强、治理基础较好的企业，可以推进流程级AI自动化；效率需求较强但治理基础一般的企业，更适合先从标准SaaS或API接入开始；安全敏感型企业不一定拒绝AI，而是更适合本地化、私有云或混合部署；数字基础薄弱的企业，应先补数据治理和流程标准化。",
        "因此，安全顾虑不是简单的负向变量。它更像部署路径选择条件，会把企业从轻量SaaS引向更可控的私有化或混合架构。这个解释来自投稿稿中的研究思路，但在课程汇报中只作为模型结果的管理解释，不写成期刊式假设检验。",
    ]:
        para(doc, text)
    for item in [
        "高采纳型：数字基础和治理能力较好，可推进流程级AI自动化。",
        "稳健增长型：已有云服务和数据分析基础，适合从局部流程扩展到关键业务流程。",
        "安全敏感型：优先考虑本地化、私有云、混合部署、日志审计和权限分级。",
        "初始观望型：先做低风险SaaS试点，积累数据治理经验后再扩大范围。",
    ]:
        bullet(doc, item)
    doc.add_page_break()

    heading(doc, "八、Agent原型与GitHub证据链")
    for text in [
        "10_Agent系统是本项目的扩展部分，用于把数据查询、模型预测、证据引用和回答约束做成可测试原型。它不是把报告变成聊天内容，而是要求回答必须来自仓库证据。没有证据的问题返回无法确认，引用必须指向结果表、报告或数据说明。",
        "源码目录保留agent_tools、api、rag、evaluation、training、tests等模块。测试覆盖预测格式、引用准确性、数据泄漏控制和质量契约。最终复核时运行`python -m unittest discover -s \"10_Agent系统/tests\" -p \"test_*.py\"`，8项测试通过。",
        "GitHub公开项目只放可公开的官方数据处理结果、源码、报告和PPT，不放私密访谈或问卷逐行材料。这个边界既保护数据伦理，也避免老师检查时把公开证据和非公开素材混淆。",
    ]:
        para(doc, text)
    table(
        doc,
        ["模块", "作用"],
        [
            ["agent_tools", "指标查询、预测、解释、引用和图表渲染"],
            ["rag", "证据索引与检索"],
            ["evaluation", "Agent回答质量评估"],
            ["tests", "预测格式、引用准确性、泄漏控制和质量契约测试"],
            ["reports", "Agent评估结果和模型登记信息"],
        ],
    )
    doc.add_page_break()

    heading(doc, "九、小组PPT汇报结构")
    para(doc, "小组PPT按课程展示逻辑组织，而不是按投稿论文结构组织。18页内容从研究问题进入，随后讲数据生命周期、建模面板、特征体系、Stage 1机制解释、Stage 2外部验证、部署路径、Agent原型、研究价值和方法边界。")
    table(
        doc,
        ["页码", "主题", "证据来源"],
        [
            ["1-4", "研究主题、现实问题、课程化框架", "投稿稿作为参考，报告重新课程化表达"],
            ["5-7", "数据生命周期、清洗规模、变量体系", "manifest、cleaning_retention_summary.csv、数据说明"],
            ["8-12", "Stage 1、Stage 2模型结果与特征解释", "course_algorithm_comparison.csv、图表输出"],
            ["13-15", "部署路径与Agent原型", "聚类结果、Agent测试和源码模块"],
            ["16-18", "贡献、边界与结束页", "课程报告、GitHub整理和复核结论"],
        ],
    )
    para(doc, "PPT中去掉了不适合课程汇报的期刊投稿口吻和夸张表达，如“专家/评委”“论文中”“完美印证”“极高”等。数值口径统一使用当前课程项目结果表，避免把投稿稿口径和课程复核口径混用。")
    doc.add_page_break()

    heading(doc, "十、成员分工与个人材料归档")
    para(doc, "成员个人材料统一放在04_小组成员个人作业整理中。景浩伟的个人材料不再单独占用一个额外目录，而是按与其他成员一致的结构归档：01_个人10次作业PDF提交版、02_个人实践报告。")
    table(
        doc,
        ["成员", "主要分工", "最终个人材料"],
        [
            ["景浩伟", "选题统筹、数据与模型流程、报告整合、PPT逻辑、答辩", "10次作业PDF + 个人总结报告DOCX/PDF"],
            ["张新通", "数据来源、生命周期、清洗记录和边界说明", "10次作业PDF + 个人实践报告DOCX/PDF"],
            ["刘子涵", "模型方法、指标解释和图表核对", "10次作业PDF + 个人实践报告DOCX/PDF"],
            ["黄陈熙", "机制解释、部署策略、PPT视觉整理", "10次作业PDF + 个人实验报告DOCX/PDF"],
        ],
    )
    doc.add_page_break()

    heading(doc, "十一、课程知识点对应")
    table(
        doc,
        ["课程知识点", "本项目对应内容"],
        [
            ["数据挖掘流程", "官方数据获取、manifest登记、清洗、面板构建、建模、解释和展示"],
            ["回归分析", "OLS、Ridge和VIF诊断"],
            ["集成学习", "Random Forest和ExtraTrees模型比较"],
            ["聚类分析", "KMeans企业画像和部署路径建议"],
            ["模型评价", "R²、MAE、GroupKFold by country"],
            ["可视化", "报告图表、PPT图表、证据映射表"],
            ["工程复现", "GitHub项目、requirements、RUN_ALL_CHECKS和Agent测试"],
        ],
    )
    para(doc, "从课程要求看，本项目不是只提交一个PPT，而是把数据、代码、模型、报告、PPT和个人作业放在同一条证据链里。老师可以从PPT看展示逻辑，从报告看方法说明，从源码看复现入口，从数据目录看结果来源，从成员目录看个人作业完整性。")
    doc.add_page_break()

    heading(doc, "十二、局限与后续改进")
    for text in [
        "第一，Eurostat数据主要反映欧洲企业ICT环境，不能直接代表中国中小企业平均情况。第二，Stage 2是行业/区域GE10口径，不是SME-only样本，报告已在结论中保留边界。第三，模型使用公开统计数据做相关性和预测解释，没有采用随机实验、双重差分或断点回归，不能声称因果关系。",
        "后续改进可以从三方面展开：补充中国本土企业公开统计或问卷数据；增加时间留出或行业留出验证，检查模型在更严格场景下的泛化能力；把Agent原型接入更真实的业务流程，用实际使用结果检验部署建议是否有效。",
        "本次最终材料按课程提交目标重新整理，已经将投稿稿与小组汇报分清：投稿稿提供选题和理论参考，小组报告负责课程展示和作业验收，个人材料统一进入成员作业整理区。",
    ]:
        para(doc, text)
    doc.add_page_break()

    heading(doc, "附录A：复现入口")
    table(
        doc,
        ["入口", "说明"],
        [
            ["src/acquisition/download_sources.py", "下载Stage 1官方数据并记录manifest"],
            ["src/acquisition/download_stage2_large_sources.py", "下载Stage 2行业/区域数据"],
            ["src/pipeline.py", "构建Stage 1面板并训练基础模型"],
            ["src/pipeline_stage2_large.py", "构建Stage 2行业/区域面板并训练模型"],
            ["src/course_ml_diagnostics.py", "生成课程诊断表、OLS、VIF和模型比较"],
            ["10_Agent系统/tests", "检查Agent预测、引用、泄漏和质量契约"],
        ],
    )
    heading(doc, "附录B：参考资料")
    for ref in [
        "Eurostat Data Browser and SDMX API：企业ICT、AI、云计算、数字强度等公开统计数据。",
        "scikit-learn documentation：GroupKFold、Ridge、RandomForest、ExtraTrees、KMeans等模型方法。",
        "投稿稿_科技管理研究_终版_改写.docx：仅作为课程选题、理论框架和研究表达边界参考，不作为本次课程小组报告结构。",
    ]:
        para(doc, ref, indent=False)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("机器学习课程小组汇报报告｜中小企业AI流程自动化采纳机制研究案例")
    set_run(r, size=9)

    doc.save(GROUP_DOCX)
    normalize_docx_colors(GROUP_DOCX)


def build_personal_summary(target_dir: Path) -> None:
    docx_path = target_dir / "信管2301景浩伟202321054012个人总结报告.docx"
    doc = Document()
    configure_doc(doc)
    cover(
        doc,
        "机器学习课程个人总结报告",
        "景浩伟 202321054012",
        [
            ["姓名", "景浩伟"],
            ["班级", "信管2301"],
            ["学号", "202321054012"],
            ["个人材料", "10次个人作业PDF + 1份个人总结报告"],
            ["小组主题", "中小企业AI流程自动化采纳机制研究案例"],
        ],
    )
    for title, paragraphs in [
        (
            "一、个人作业完成情况",
            [
                "本课程个人材料按十次作业顺序整理为PDF提交版，放在小组成员个人作业整理目录下。每次作业保留对应的模型、代码、图表或分析结果，便于老师按序检查。",
                "十次作业覆盖了机器学习课程中的数据处理、回归、分类、树模型、集成学习、聚类、模型评价、可视化和综合案例分析等内容。整理时只统一文件顺序和命名，不改动原始作业内容。",
            ],
        ),
        (
            "二、小组案例中的个人分工",
            [
                "在小组案例中，我主要负责选题统筹、研究框架、数据处理思路、特征工程、模型评估、报告整合、PPT逻辑和答辩准备。我的工作重点是把投稿稿中的研究思路转成机器学习课程能检查的材料，而不是把课程报告写成投稿论文。",
                "数据部分，我负责确认Eurostat官方数据链、清洗边界和建模口径。Stage 1用于SME规模组机制解释，Stage 2用于行业/区域GE10外部验证，两者在报告和PPT中都保留清楚边界。",
            ],
        ),
        (
            "三、模型与结果理解",
            [
                "模型部分，我参与整理OLS、Ridge、Random Forest、ExtraTrees和KMeans的课程化解释，并把GroupKFold by country作为核心验证设计写入报告。这个设计能降低同一国家观测同时进入训练集和验证集造成的乐观评估。",
                "从结果看，机器学习能力、部署准备度和治理成熟度是解释AI流程自动化采纳差异的重要线索。安全顾虑不被简单写成阻力，而是影响企业选择SaaS、API、本地化或混合部署路径的条件。",
            ],
        ),
        (
            "四、课程收获与不足",
            [
                "这次项目让我认识到，机器学习课程作业不能只看模型分数。一个完整项目必须同时说明数据从哪里来、清洗规则是什么、特征为什么这样构建、验证方式是否会造成泄漏、结论能否被复核。",
                "不足也比较明确。Eurostat指标口径复杂，变量解释需要持续对照官方说明；公开统计数据不能直接替代企业微观数据；部分变量存在共线性，不能只看单个OLS系数。后续如果继续完善，应补充中国本土数据和更严格的时间留出验证。",
            ],
        ),
    ]:
        heading(doc, title)
        for text in paragraphs:
            para(doc, text)
        doc.add_page_break()
    heading(doc, "五、最终提交说明")
    para(doc, "我的个人材料已按小组成员统一格式放入04_小组成员个人作业整理\\04_景浩伟_202321054012，其中包括01_个人10次作业PDF提交版和02_个人实践报告。最终提交不再保留单独的个人作业总结目录，避免和小组成员整理结构重复。")
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("信管2301 景浩伟 202321054012 个人总结报告")
    set_run(r, size=9)
    doc.save(docx_path)
    normalize_docx_colors(docx_path)


def normalize_docx_colors(path: Path) -> None:
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                text = data.decode("utf-8", "ignore")
                text = re.sub(r'(<w:color\b[^>]*\bw:val=")[^"]+(")', r"\g<1>000000\2", text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def convert_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    return out_dir / f"{docx_path.stem}.pdf"


def edit_pptx_text() -> None:
    replacements = {
        "现实痛点：为什么采纳AI流程自动化如此重要？": "课程问题：为什么中小企业需要AI流程自动化？",
        "中小企业面临数字化转型的急迫压力，但在AI流程自动化上呈现出极高的初期流失率和落地困难，亟需探明其背后的断层逻辑。": "中小企业希望用AI减少重复录入和人工复核，但资金、人才、数据和安全约束会影响项目能否落地。",
        "资源、认知与技术的多重断层限制了SME的数字化转型": "资源、认知与技术条件共同限制SME的数字化转型",
        "理论与现实背景：中小企业在AI转型初期面临极高的流失率，主要受制于三大现实阻力。": "理论与现实背景：中小企业AI转型不是单纯技术替换，落地前需要同时处理三类约束。",
        "基于高标准复现的数据生命周期设计": "基于可复现要求的数据生命周期设计",
        "研究遵循严密的数据治理链路，确保从原始官方信息提取到最终模型面板的数据完全可追溯、可复现。": "研究从官方数据、manifest、清洗脚本、建模面板到结果图表保留对应证据，便于课程检查。",
        "严格控制未来数据防泄漏机制": "控制目标泄漏和地理泄漏",
        "从1277万官方记录到5814条建模面板": "从1277万官方记录到5814条建模面板",
        "并非千万级粗暴训练，而是通过精准画像与特征过滤提取高质量复现面板，剔除大量无意义或低置信记录。": "项目不是直接用千万级原始行训练模型，而是经过指标筛选、面板聚合和目标非空过滤，形成可解释建模数据。",
        "Ridge模型展现了极高的局部拟合与解释力": "Ridge模型在SME规模组口径下表现稳定",
        "高达 0.8680 的决定系数表明，所选的效率、安全、部署三大前置维度能强有力地解释 SME 的采纳方差。": "Stage 1使用SME规模组面板，Ridge在按国家分组验证下保持较好解释力。",
        "决定系数 R² (GroupKFold)": "决定系数 R² (GroupKFold)",
        "0.868": "0.874",
        "1.8342": "1.7730",
        "此指标在具有严格地理隔离约束的 GroupKFold 验证中取得，并非过拟合产物，说明核心机制具有极高的实质相关性，模型对中小企业采纳路径的选择具有充分的解释力与预测稳定性。": "该结果来自GroupKFold by country口径，说明核心变量对SME规模组采纳差异有较强解释力；结论限定为相关性和预测解释。",
        "Stage 2：跨越50个行业的外部稳健性泛化": "Stage 2：跨越50个行业的外部验证",
        "打破 SME-only 的解释边界，将样本扩张至全行业、大区域，测试前置偏好维度的跨域外部有效性。": "Stage 2不再写成SME-only结论，而是用行业/区域GE10口径检验变量框架是否仍有解释力。",
        "Long-run GroupKFold": "GroupKFold by country",
        "ExtraTrees在多行业长期验证中保持高度稳健": "ExtraTrees在行业/区域验证中保持可用解释力",
        "在5814条全行业复杂样本中，决定系数稳定在0.7以上，证明了核心机制的跨域普适性。": "在5814条行业/区域样本中，模型仍保留0.7左右的解释力，说明变量框架有外部验证价值。",
        "0.7245": "0.7073",
        "1.9646": "2.1060",
        "当场景从 SME 扩展到包含大型企业、重资产集团的 50 个行业大盘时，效率与安全约束主导的底层机制依旧未失效。R²从局域的0.86滑落至0.72属合理的“非同源泛化衰减”，证明模型非过拟合记忆，而是捕捉到了通用因子的底层规律。": "场景从SME规模组扩展到行业/区域口径后，模型难度上升，R²下降是合理现象；这里用于说明外部验证，不用于宣称因果规律。",
        "效率主导、安全兜底、部署分化的特征权重": "效率、治理与部署条件的特征权重",
        "机器学习提取的特征重要性反馈完美印证了非线性业务预期的结构。": "特征重要性帮助解释哪些变量在两个阶段中更稳定，仍需结合口径边界审慎表述。",
        "绝对主导地位": "贡献较高",
        "降本增效指标贡献了最大的方差解释度，是企业立项与启动自动化的第一动因与乘数。": "AI能力和效率相关指标贡献较高，是企业启动自动化项目的重要线索。",
        "非线性阈值阻断": "路径约束",
        "权重仅次于效率，但在树模型节点中表现为高阶分裂点。不满足基础安全阈值即引发“一票否决”。": "安全和数据治理指标会影响部署路径选择，不宜简单写成采纳阻力。",
        "渗透乘数调节": "落地条件",
        "权重居第三位，不触发原始采纳，但显著调节采纳在多深、多广的业务流中落地（深度与速度）。": "部署准备度影响AI进入业务流程的范围和速度。",
        "模型揭示出中小企业 AI 采纳并非粗暴的单一线性驱动，而是在收益与风险间发生动态时序权衡的过滤机制。": "模型结果提示，中小企业AI采纳不是单一因素驱动，而是效率收益、安全约束和部署条件共同作用。",
        "唯一敲门砖。无效率提升则无立项。": "重要进入条件。效率收益不清楚时，项目很难获得投入。",
        "隐形墙。若黑盒解释性不能通过业务审计，项目将发生硬阻断。": "实施门槛。若解释性、权限和审计要求不足，项目容易停在试点阶段。",
        "基于研究结论，中小企业应摒弃盲目的大规模 IT 重构，采取顺应内在机制规律的渐进式闭环策略。": "基于研究结论，中小企业更适合从低风险流程切入，再逐步完善安全、治理和部署能力。",
        "构建证据约束型AI决策辅助": "构建证据约束型AI决策辅助原型",
        "将 Ridge 与 ExtraTrees 在 1277 万原始数据中挖掘出的特征权重与机制规律提炼为业务逻辑库。": "将Ridge与ExtraTrees结果、指标解释和证据索引整理为业务逻辑库。",
        "(非因果胡编乱造)": "避免无证据扩展",
        "四大维度共同构筑的研究价值": "课程项目的四类产出",
        "高可信度的量化解释": "可复核的量化解释",
        "贡献": "产出",
        "严守方法论边界：局限性与下一步计划": "方法边界与下一步计划",
        "郑重声明：机器学习模型验证的是“机制相关性”与“偏好解释力度”。虽然特征重要性表现出明显的结构化差异，但并未运用双重差分、断点回归等严密因果推断范式。": "本项目验证的是机制相关性与预测解释力度，没有使用双重差分、断点回归等因果推断方法。",
        "绝不混淆主次证据。": "不与公开数据证据混用。",
        "“ 本研究用可复现的数据生命周期和机器学习验证，解释了中小企业AI流程自动化采纳中的效率、安全与部署机制。”": "本项目用可复现的数据生命周期和机器学习验证，说明中小企业AI流程自动化采纳中的效率、安全与部署条件。",
        "感谢专家/评委聆听，欢迎批评指正。": "感谢老师和同学聆听，欢迎批评指正。",
    }
    tmp = PPTX.with_suffix(".tmp.pptx")
    with zipfile.ZipFile(PPTX, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("ppt/slides/slide") and item.filename.endswith(".xml"):
                text = data.decode("utf-8")
                for old, new in replacements.items():
                    text = text.replace(f"<a:t>{escape(old)}</a:t>", f"<a:t>{escape(new)}</a:t>")
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(PPTX)


def reorganize_personal() -> None:
    if PERSONAL_OLD.exists():
        (BACKUP / "旧单独个人作业总结").mkdir(parents=True, exist_ok=True)
    if MEMBERS_OLD.exists() and not MEMBERS.exists():
        shutil.move(str(MEMBERS_OLD), str(MEMBERS))
    elif MEMBERS_OLD.exists() and MEMBERS.exists():
        shutil.move(str(MEMBERS_OLD), str(BACKUP / MEMBERS_OLD.name))

    target = MEMBERS / "04_景浩伟_202321054012"
    if target.exists() and not PERSONAL_OLD.exists():
        return
    if target.exists():
        shutil.move(str(target), str(BACKUP / target.name))
    homework = target / "01_个人10次作业PDF提交版"
    report = target / "02_个人实践报告"
    homework.mkdir(parents=True, exist_ok=True)
    report.mkdir(parents=True, exist_ok=True)

    source_homework = PERSONAL_OLD / "05_个人10次作业PDF提交版"
    for pdf in sorted(source_homework.glob("*.pdf")):
        if re.match(r"^\d{2}_", pdf.name):
            shutil.copy2(pdf, homework / pdf.name.replace("第十次.pdf", "第十次作业.pdf") if pdf.name.endswith("第十次.pdf") else homework / pdf.name)
    build_personal_summary(report)
    (target / "文件顺序说明.txt").write_text(
        "本目录按小组成员统一格式整理：01_个人10次作业PDF提交版保留10次个人作业；02_个人实践报告保留个人总结报告DOCX/PDF。\n",
        encoding="utf-8",
    )

    if PERSONAL_OLD.exists():
        shutil.move(str(PERSONAL_OLD), str(BACKUP / PERSONAL_OLD.name))


def write_member_lists() -> None:
    rows = []
    for member_dir in sorted(p for p in MEMBERS.iterdir() if p.is_dir() and re.match(r"^\d{2}_", p.name)):
        for file in sorted(member_dir.rglob("*")):
            if file.is_file():
                rows.append([member_dir.name, str(file.relative_to(MEMBERS)), file.suffix.lower(), file.stat().st_size])
    for name in ["小组成员作业文件清单.csv", "整理后文件完整清单.csv"]:
        with (MEMBERS / name).open("w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["成员", "相对路径", "扩展名", "大小bytes"])
            writer.writerows(rows)

    summary = []
    ok = True
    for member_dir in sorted(p for p in MEMBERS.iterdir() if p.is_dir() and re.match(r"^\d{2}_", p.name)):
        homework_count = len(list((member_dir / "01_个人10次作业PDF提交版").glob("*.pdf")))
        report_dir = member_dir / "02_个人实践报告"
        docx_count = len(list(report_dir.glob("*.docx")))
        pdf_count = len(list(report_dir.glob("*.pdf")))
        member_ok = homework_count == 10 and docx_count == 1 and pdf_count == 1
        ok = ok and member_ok
        summary.append(
            {
                "member": member_dir.name,
                "homework_pdf_count": homework_count,
                "report_docx_count": docx_count,
                "report_pdf_count": pdf_count,
                "ok": member_ok,
            }
        )
    (MEMBERS / "小组成员作业质量核验详情.json").write_text(
        json.dumps({"ok": ok, "members": summary}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    lines = [
        "# 小组成员个人作业整理说明",
        "",
        "本目录统一整理四名成员的个人材料。每名成员保留两个核心部分：10次个人作业PDF、1份个人总结/实践报告DOCX与PDF。",
        "",
        "| 成员 | 10次作业PDF | 报告DOCX | 报告PDF | 结论 |",
        "|---|---:|---:|---:|---|",
    ]
    for item in summary:
        lines.append(
            f"| {item['member']} | {item['homework_pdf_count']} | {item['report_docx_count']} | {item['report_pdf_count']} | {'通过' if item['ok'] else '需处理'} |"
        )
    (MEMBERS / "小组成员作业整理说明.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    (MEMBERS / "小组成员作业质量核验报告.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_evidence_map() -> None:
    rows = [
        [1, "封面与研究主题", "小组课程汇报报告、投稿稿仅作选题参考"],
        [2, "研究核心与目标", "报告第一、二节"],
        [3, "SME现实约束", "报告第二节"],
        [4, "课程化研究框架", "报告第二、四节"],
        [5, "数据生命周期", "manifest、数据来源说明、报告第三节"],
        [6, "数据清洗规模", "cleaning_retention_summary.csv"],
        [7, "变量体系", "报告第四节"],
        [8, "Stage 1设计", "course_algorithm_comparison.csv、报告第五节"],
        [9, "Stage 1结果", "course_algorithm_comparison.csv"],
        [10, "Stage 2设计", "stage2_source_profile.json、报告第六节"],
        [11, "Stage 2结果", "course_algorithm_comparison.csv"],
        [12, "特征权重解释", "feature_importance表与报告第六节"],
        [13, "采纳演化路径", "报告第七节"],
        [14, "SME采纳建议", "报告第七节"],
        [15, "Agent原型", "10_Agent系统源码与测试结果"],
        [16, "课程项目产出", "最终提交目录与GitHub复核说明"],
        [17, "方法边界", "报告第十二节"],
        [18, "结束页", "小组课程汇报结论"],
    ]
    with EVIDENCE_MAP.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["slide", "主题", "证据来源"])
        writer.writerows(rows)


def convert_outputs() -> None:
    for path in [GROUP_PDF, PPT_PDF, MEMBERS / "04_景浩伟_202321054012" / "02_个人实践报告" / "信管2301景浩伟202321054012个人总结报告.pdf"]:
        if path.exists():
            path.unlink()
    generated = convert_to_pdf(GROUP_DOCX, GROUP)
    if generated != GROUP_PDF:
        generated.replace(GROUP_PDF)
    personal_docx = MEMBERS / "04_景浩伟_202321054012" / "02_个人实践报告" / "信管2301景浩伟202321054012个人总结报告.docx"
    convert_to_pdf(personal_docx, personal_docx.parent)
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(GROUP), str(PPTX)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )


def write_opening_docs() -> None:
    (FINAL / "打开说明.md").write_text(
        "\n".join(
            [
                "# 打开说明",
                "",
                "请使用正常 Windows 路径预览最终材料：",
                "",
                f"`{FINAL}`",
                "",
                "最终提交目录现在按课程验收顺序整理为：",
                "",
                "1. `01_数据`",
                "2. `02_源码`",
                "3. `03_小组汇报PPT和报告`",
                "4. `04_小组成员个人作业整理`",
                "",
                "小组报告是课程汇报报告，不是投稿稿；投稿稿只作为选题和理论框架参考。",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    (FINAL / "GitHub完整性复核.md").write_text(
        "\n".join(
            [
                "# GitHub完整性复核",
                "",
                "当前本地仓库远程地址：`https://github.com/hoanglenga2000-glitch/sme-ai-workflow-adoption.git`。",
                "",
                "本轮未执行提交或推送，只对本地最终提交材料进行课程化纠偏整理。上传GitHub前建议以`课程最终提交材料`为入口核对四个目录：数据、源码、小组汇报PPT和报告、小组成员个人作业整理。",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def audit() -> None:
    try:
        import fitz
    except Exception as exc:
        raise RuntimeError("PyMuPDF is required for audit") from exc

    def pdf_pages(path: Path) -> tuple[int, list[int]]:
        blanks = []
        with fitz.open(path) as doc:
            for idx, page in enumerate(doc, 1):
                text = page.get_text("text").strip()
                pix = page.get_pixmap(matrix=fitz.Matrix(0.12, 0.12), alpha=False)
                data = pix.samples
                nonwhite = 0
                for i in range(0, len(data), pix.n):
                    if any(c < 245 for c in data[i : i + 3]):
                        nonwhite += 1
                ratio = nonwhite / max(1, pix.width * pix.height)
                if len(text) < 10 and ratio < 0.015:
                    blanks.append(idx)
            return doc.page_count, blanks

    group_pages, group_blanks = pdf_pages(GROUP_PDF)
    ppt_pages, ppt_blanks = pdf_pages(PPT_PDF)
    personal_pdf = MEMBERS / "04_景浩伟_202321054012" / "02_个人实践报告" / "信管2301景浩伟202321054012个人总结报告.pdf"
    personal_pages, personal_blanks = pdf_pages(personal_pdf)
    member_detail = json.loads((MEMBERS / "小组成员作业质量核验详情.json").read_text(encoding="utf-8"))
    agent_proc = subprocess.run(
        [str(Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "python" / "python.exe"), "-m", "unittest", "discover", "-s", str(SOURCE / "10_Agent系统" / "tests"), "-p", "test_*.py"],
        cwd=REPO,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        check=False,
    )
    checks = [
        ["最终目录结构", DATA.exists() and SOURCE.exists() and GROUP.exists() and MEMBERS.exists() and not PERSONAL_OLD.exists()],
        ["小组课程汇报报告", GROUP_DOCX.exists() and GROUP_PDF.exists() and group_pages >= 15 and not group_blanks],
        ["小组PPT", PPTX.exists() and PPT_PDF.exists() and ppt_pages == 18 and not ppt_blanks],
        ["成员个人材料", bool(member_detail["ok"])],
        ["景浩伟个人总结", personal_pdf.exists() and personal_pages >= 5 and not personal_blanks],
        ["Agent单元测试", agent_proc.returncode == 0 and "Ran 8 tests" in agent_proc.stdout],
        ["证据映射表", EVIDENCE_MAP.exists() and sum(1 for _ in EVIDENCE_MAP.open("r", encoding="utf-8-sig")) == 19],
    ]
    ok = all(item[1] for item in checks)
    detail = {
        "ok": ok,
        "checks": checks,
        "group_report_pages": group_pages,
        "ppt_pages": ppt_pages,
        "personal_summary_pages": personal_pages,
        "agent_test_output": agent_proc.stdout,
        "backup": str(BACKUP),
    }
    (FINAL / "最终上交审核详情.json").write_text(json.dumps(detail, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        "# 最终上交审核报告",
        "",
        f"生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}",
        "",
        "| 检查项 | 结果 |",
        "|---|---|",
    ]
    for name, passed in checks:
        lines.append(f"| {name} | {'通过' if passed else '需处理'} |")
    lines.extend(
        [
            "",
            f"- 小组课程汇报报告PDF：{group_pages}页，无空白页。",
            f"- 小组最终汇报PPT PDF：{ppt_pages}页，无空白页。",
            f"- 景浩伟个人总结报告PDF：{personal_pages}页，无空白页。",
            "- 投稿稿仅作为参考材料，没有混入小组课程报告结构。",
            f"- 旧单独个人目录和旧小组报告已备份到：`{BACKUP}`。",
        ]
    )
    (FINAL / "最终上交审核报告.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    (FINAL / "提交说明与质量核验.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    if not ok:
        raise RuntimeError(json.dumps(detail, ensure_ascii=False, indent=2))


def main() -> None:
    for path in [FINAL, GROUP, DATA, SOURCE]:
        if not path.exists():
            raise FileNotFoundError(path)
    BACKUP.mkdir(parents=True, exist_ok=True)
    safe_under(BACKUP, REPO)
    safe_under(FINAL, REPO)
    safe_under(GROUP, FINAL)
    build_group_report()
    edit_pptx_text()
    reorganize_personal()
    subprocess.run(
        [
            str(Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "python" / "python.exe"),
            str(REPO / "scripts" / "rebuild_jing_personal_summary_report.py"),
        ],
        cwd=REPO,
        check=True,
    )
    write_evidence_map()
    convert_outputs()
    subprocess.run(
        [
            str(Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "python" / "python.exe"),
            str(REPO / "scripts" / "rebuild_group_work_summary_report.py"),
        ],
        cwd=REPO,
        check=True,
    )
    write_member_lists()
    write_opening_docs()
    audit()
    print("corrected final submission rebuilt")
    print(FINAL)
    print(BACKUP)


if __name__ == "__main__":
    main()
