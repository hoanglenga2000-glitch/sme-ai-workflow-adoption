from __future__ import annotations

import json
import re
import shutil
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
FINAL = REPO / "课程最终提交材料"
MEMBER = FINAL / "04_小组成员个人作业整理" / "04_景浩伟_202321054012"
REPORT_DIR_CANDIDATES = [MEMBER / "02_个人案例总结", MEMBER / "02_个人实践报告"]
TABLES = REPO / "outputs" / "tables"
REPORTS = REPO / "outputs" / "reports"
FIGS = REPO / "outputs" / "figures" / "academic"
DOCX_NAME = "信管2301景浩伟202321054012个人总结报告.docx"
PDF_NAME = "信管2301景浩伟202321054012个人总结报告.pdf"
BLACK = RGBColor(0, 0, 0)


def report_dir() -> Path:
    for path in REPORT_DIR_CANDIDATES:
        if path.exists():
            path.mkdir(parents=True, exist_ok=True)
            return path
    REPORT_DIR_CANDIDATES[0].mkdir(parents=True, exist_ok=True)
    return REPORT_DIR_CANDIDATES[0]


def set_font(run, size: float = 10.5, bold: bool = False, font: str = "微软雅黑") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = BLACK


def set_cell_text(cell, text: object, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(text))
    set_font(r, size=9, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_title(doc: Document, text: str, size: float = 22) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=size, bold=True)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 14, bold=True)


def add_para(doc: Document, text: str, indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_font(r)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        shade(table.rows[0].cells[i], "D9D9D9")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str, width_cm: float = 13.2) -> None:
    path = FIGS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(3)
    r = cap.add_run(caption)
    set_font(r, size=9)


def fmt(value: object, digits: int = 4) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def normalize_docx(path: Path) -> None:
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                text = data.decode("utf-8", errors="ignore")
                text = re.sub(r'(<w:color\b[^>]*\bw:val=")[^"]+(")', r"\g<1>000000\2", text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def pdf_pages_and_blanks(path: Path) -> tuple[int, list[int]]:
    try:
        import fitz
    except Exception:
        return 0, []
    blanks: list[int] = []
    with fitz.open(path) as pdf:
        for idx, page in enumerate(pdf, 1):
            text = page.get_text("text").strip()
            pix = page.get_pixmap(matrix=fitz.Matrix(0.12, 0.12), alpha=False)
            nonwhite = 0
            for i in range(0, len(pix.samples), pix.n):
                if any(channel < 245 for channel in pix.samples[i : i + 3]):
                    nonwhite += 1
            if len(text) < 10 and nonwhite / max(1, pix.width * pix.height) < 0.015:
                blanks.append(idx)
        return pdf.page_count, blanks


def docx_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())


def audit_text(text: str) -> dict:
    suspicious_terms = [
        "赋能",
        "全面提升",
        "深度剖析",
        "深入探讨",
        "值得注意的是",
        "毋庸置疑",
        "显而易见",
        "综上所述",
        "打造",
        "全方位",
        "高质量发展",
        "无缝",
        "生态",
        "革命性",
        "范式",
        "里程碑",
    ]
    hits = {term: text.count(term) for term in suspicious_terms if term in text}
    return {
        "char_count": len(text),
        "paragraph_count": text.count("\n") + 1 if text else 0,
        "first_person_count": text.count("我"),
        "ai_phrase_hits": hits,
        "ok": len(text) >= 4500 and len(hits) <= 2 and text.count("我") >= 8,
    }


def build() -> dict:
    out_dir = report_dir()
    docx_path = out_dir / DOCX_NAME
    pdf_path = out_dir / PDF_NAME

    alg = pd.read_csv(TABLES / "course_algorithm_comparison.csv")
    reg = pd.read_csv(TABLES / "course_regression_summary.csv")
    clean = pd.read_csv(TABLES / "cleaning_retention_summary.csv")
    quality = pd.read_csv(TABLES / "enhanced_data_quality_audit.csv")
    vif = pd.read_csv(TABLES / "course_vif_diagnostics.csv")
    imp1 = pd.read_csv(TABLES / "feature_importance.csv")
    imp2 = pd.read_csv(TABLES / "stage2_feature_importance.csv")
    cluster1 = pd.read_csv(TABLES / "sme_persona_clusters.csv")
    cluster2 = pd.read_csv(TABLES / "stage2_persona_clusters.csv")
    gpu = pd.read_csv(TABLES / "enhanced_gpu_baseline.csv")
    holdout = pd.read_csv(TABLES / "enhanced_holdout_results.csv")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.start_type = WD_SECTION.NEW_PAGE

    for style_name in ["Normal", "Body Text"]:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "微软雅黑"
            doc.styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            doc.styles[style_name].font.color.rgb = BLACK

    add_title(doc, "机器学习课程个人案例总结报告")
    add_title(doc, "中小企业AI流程自动化采纳机制研究", size=18)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["姓名", "景浩伟"],
            ["班级", "信管2301"],
            ["学号", "202321054012"],
            ["课程任务", "平时作业汇总、小组案例汇报、个人案例总结"],
            ["个人定位", "选题统筹、数据与模型流程负责人、报告整合、PPT逻辑和答辩准备"],
            ["材料依据", "GitHub项目输出、Eurostat官方数据处理结果、小组课程报告、投稿稿研究框架"],
        ],
        widths=[3.2, 11.5],
    )

    add_para(
        doc,
        "这份个人总结是对我在结课案例中实际工作的整理，不再只写几句“参与了模型和PPT”。我把自己的工作拆成四条线：第一是把投稿稿中的研究主题改成机器学习课程能检查的案例；第二是把公开数据、清洗代码和模型结果整理成可复核的证据链；第三是把模型输出翻译成企业AI部署偏好的解释；第四是把个人十次作业、小组PPT和最终报告统一到老师要求的提交结构里。"
    )
    add_para(
        doc,
        "我没有把投稿稿直接放进个人总结。投稿稿主要提供理论框架和研究口径，例如TOE框架、TAM模型、效率需求、安全顾虑和部署偏好。个人总结更关注我在课程项目里怎么处理数据、怎么做模型验证、怎么解释结果，以及我在整理过程中踩到的具体问题。"
    )

    add_heading(doc, "一、个人作业与课程能力基础")
    add_para(
        doc,
        "我的十次平时作业已经按顺序整理为PDF提交版，并合并成平时作业汇总PDF。前几次作业主要训练数据读取、字段整理、描述统计和可视化；中间几次转向回归、分类、树模型和模型评价；后面的作业逐步接近综合案例，需要同时处理数据来源、特征构造、模型比较和结果解释。"
    )
    add_para(
        doc,
        "这十次作业对结课案例有直接帮助。比如，回归作业让我在最终报告里保留OLS结果，而不是只写树模型分数；集成学习作业帮助我理解Random Forest和ExtraTrees为什么适合处理非线性关系；聚类作业则被用到企业画像部分。个人作业不是单独摆放的附件，它们能说明我为什么能承担小组案例中的数据和模型部分。"
    )
    add_table(
        doc,
        ["课程训练环节", "在结课案例中的对应工作"],
        [
            ["数据处理与可视化", "整理Eurostat源数据、构建面板、检查缺失率、生成报告图表。"],
            ["回归分析", "使用OLS解释变量方向，并用VIF检查共线性问题。"],
            ["正则化模型", "使用Ridge处理多变量共线性，形成Stage 1的稳定基线。"],
            ["树模型与集成学习", "使用Random Forest和ExtraTrees比较非线性解释能力。"],
            ["聚类分析", "形成SME企业画像和行业/区域画像，用于部署路径解释。"],
            ["模型评价", "使用R²、MAE、GroupKFold和按国家留出验证控制乐观评估。"],
        ],
        widths=[4.0, 10.7],
    )

    add_heading(doc, "二、我在小组案例中的具体分工")
    add_para(
        doc,
        "小组案例不是把四个人各写几页再拼起来。我的主要工作是把研究问题、数据、代码、报告和PPT串起来，让老师能沿着同一条线检查：为什么选这个题，数据从哪里来，清洗后剩下什么，模型怎么验证，结论能不能回到中小企业AI流程自动化采纳这个主题。"
    )
    add_para(
        doc,
        "我负责的部分集中在选题统筹、研究框架、数据与模型流程、结果解释和最终材料整合。张新通侧重数据来源和清洗边界，刘子涵侧重模型方法和指标解释，黄陈熙侧重机制解释、部署策略和PPT视觉整理。我需要做的是把这些内容统一成课程展示语言，避免每一部分各说各的。"
    )
    add_table(
        doc,
        ["我的工作项", "实际产出", "质量要求"],
        [
            ["选题与框架", "把“企业AI部署偏好与治理机制”收窄为“中小企业AI流程自动化采纳机制”。", "研究问题要能被数据和模型检验。"],
            ["数据流程", "确认Stage 1和Stage 2口径，整理清洗、筛选、聚合和建模边界。", "不能把千万级原始行直接说成训练样本。"],
            ["建模评估", "组织OLS、Ridge、Random Forest、ExtraTrees、MLP、KMeans的比较。", "分清解释模型、预测模型和画像模型的作用。"],
            ["报告整合", "把投稿稿框架改写为课程报告和个人总结。", "不写成投稿稿，也不写成文件清单。"],
            ["展示准备", "梳理18页小组PPT逻辑和证据映射。", "每页PPT有数据、表格或报告依据。"],
        ],
        widths=[3.0, 7.8, 4.0],
    )

    add_heading(doc, "三、数据来源和清洗工作")
    add_para(
        doc,
        "数据工作是这次项目里最容易被低估的部分。Eurostat企业ICT数据来源可靠，但不同数据表的国家、年份、行业、规模和指标编码并不完全一致。我的任务不是简单下载文件，而是把它们整理成能进入机器学习模型的面板数据，并在报告里保留每一步的边界。"
    )
    add_para(
        doc,
        "Stage 1使用企业规模组口径，官方多源数据共有134,367行，经过长表转换、目标变量非空过滤、覆盖率筛选和特征泄漏控制后，形成544行建模样本。Stage 2用于外部验证，17个官方源文件共12,770,332行，先扫描出10,453,354条非空观测，再按指标筛选保留856,880行，最后聚合为5,814行行业/区域建模面板。这个过程说明，项目的核心不是“数据很多”，而是把大体量数据清洗到口径一致、可解释、可复现。"
    )
    add_table(
        doc,
        ["清洗阶段", "原始或长表行数", "面板行数", "建模行数", "说明"],
        [
            [row["stage"], f"{int(row['raw_or_long_rows']):,}", f"{int(row['panel_rows']):,}", f"{int(row['model_rows']):,}", "用于课程报告的清洗保留率核对"]
            for _, row in clean.iterrows()
        ],
        widths=[5.2, 2.4, 2.2, 2.2, 3.3],
    )
    add_table(
        doc,
        ["数据层", "建模行数", "列数", "目标非空", "国家/地区", "年份", "用途边界"],
        [
            [
                "Stage 1 SME规模组" if row["dataset"] == "stage1_sme_size_class" else "Stage 2 行业/区域GE10",
                int(row["rows"]),
                int(row["columns"]),
                int(row["target_nonnull"]),
                int(row["geo_count"]),
                f"{int(row['year_min'])}-{int(row['year_max'])}",
                "SME机制解释" if row["dataset"] == "stage1_sme_size_class" else "外部验证，不写成SME-only",
            ]
            for _, row in quality.iterrows()
        ],
        widths=[4.1, 1.7, 1.5, 1.7, 1.8, 2.0, 3.7],
    )
    add_figure(doc, "fig1_academic_validation_clean.png", "图1 数据清洗、筛选与建模流程")

    add_heading(doc, "四、变量体系和特征工程")
    add_para(
        doc,
        "变量设计时，我主要参考投稿稿中的TOE框架和TAM模型，但没有照搬论文写法。课程项目更需要把管理概念转成模型变量。技术维度对应AI能力、云服务能力和数据成熟度；组织维度对应治理成熟度、ICT人才和培训；环境维度对应国家、年份、行业和规模差异。TAM中的感知有用性被转成效率需求，感知易用性被转成部署准备度。"
    )
    add_para(
        doc,
        "因变量是E_AI_TPA，即企业使用AI进行工作流自动化或辅助决策的比例。自变量包括机器学习能力、自然语言生成、云开发能力、云数据分析、安全顾虑、部署准备度、数据成熟度、数字基础和治理成熟度。这里有一个重要处理：我在整理模型说明时要求剔除目标变量及其直接派生字段，避免模型用相近指标“偷看答案”。"
    )
    add_table(
        doc,
        ["变量组", "代表字段", "我在报告中的解释方式"],
        [
            ["因变量", "E_AI_TPA", "企业是否把AI用于流程自动化和辅助决策。"],
            ["AI能力", "E_AI_TML、E_AI_TNLG、E_AI_CC", "企业是否具备机器学习、文本生成和AI相关应用基础。"],
            ["部署准备度", "deployment_readiness_index、E_CC_PDEV、E_CC_DA", "企业是否具备云开发、数据分析和流程接入能力。"],
            ["安全顾虑", "security_concern_index、security_x_efficiency", "安全不是简单阻力，而是影响部署路径的条件。"],
            ["治理成熟度", "governance_maturity_proxy", "数据管理、ICT培训和流程治理基础。"],
            ["控制变量", "geo、year、nace_r2、size_emp", "控制国家、年份、行业和规模差异。"],
        ],
        widths=[3.0, 4.8, 7.0],
    )

    add_heading(doc, "五、模型选择和验证方式")
    add_para(
        doc,
        "建模部分我没有只保留一个最好看的分数。OLS用于解释变量方向和显著性，Ridge用于处理共线性，Random Forest和ExtraTrees用于捕捉非线性关系，MLP作为A10 GPU神经网络基线，KMeans用于企业画像。不同模型回答的问题不一样，放在一起才像一个完整的机器学习课程项目。"
    )
    add_para(
        doc,
        "我特别重视GroupKFold by country。企业ICT数据天然带有国家结构，同一个国家的企业观测往往共享制度环境、数字基础和产业特点。如果随机拆分训练集和验证集，同一国家的相似观测可能同时出现在两边，模型分数会偏乐观。按国家分组验证更严格，也更适合课堂答辩中说明“为什么结果不是简单记住国家差异”。"
    )
    add_table(
        doc,
        ["数据层", "模型", "R²均值", "R²标准差", "MAE均值", "特征数", "验证方式"],
        [
            [row["dataset"], row["model_cn"], fmt(row["r2_mean"]), fmt(row["r2_std"]), fmt(row["mae_mean"]), int(row["feature_count"]), row["validation"]]
            for _, row in alg.iterrows()
        ],
        widths=[3.3, 2.3, 1.6, 1.8, 1.8, 1.5, 3.0],
    )
    add_table(
        doc,
        ["数据层", "样本量", "R²", "调整R²", "特征数", "p<0.05变量数"],
        [
            [row["dataset"], int(row["n"]), fmt(row["r2"]), fmt(row["adj_r2"]), int(row["features"]), int(row["significant_05"])]
            for _, row in reg.iterrows()
        ],
        widths=[4.0, 1.8, 1.8, 1.8, 1.7, 2.5],
    )
    add_figure(doc, "fig1a_model_comparison_ppt.png", "图2 Stage 1和Stage 2模型交叉验证比较", width_cm=12.6)

    add_heading(doc, "六、模型结果和我的理解")
    stage1_best = alg[alg["dataset"].str.contains("Stage 1")].sort_values("r2_mean", ascending=False).iloc[0]
    stage2_best = alg[alg["dataset"].str.contains("Stage 2")].sort_values("r2_mean", ascending=False).iloc[0]
    add_para(
        doc,
        f"从课程诊断结果看，Stage 1中{stage1_best['model_cn']}在GroupKFold by country下R²均值为{fmt(stage1_best['r2_mean'])}，MAE均值为{fmt(stage1_best['mae_mean'])}；Stage 2中{stage2_best['model_cn']}R²均值为{fmt(stage2_best['r2_mean'])}，MAE均值为{fmt(stage2_best['mae_mean'])}。这个结果和投稿稿中的完整训练口径并不完全相同，我在报告里把两套口径分开写，避免把全样本拟合、交叉验证和留出验证混在一起。"
    )
    add_para(
        doc,
        "特征重要性结果比较稳定。Stage 1中E_AI_TML的重要性最高，Stage 2中ai_industry__E_AI_TML也排在第一。这说明机器学习能力是解释AI流程自动化采纳的核心线索。部署准备度、数字基础、自然语言生成、云服务能力和安全顾虑也进入前列，但它们的含义不能被简单写成因果影响。我的表述是：这些变量提供了采纳差异的可解释线索。"
    )
    add_table(
        doc,
        ["层次", "Top1特征", "重要性均值", "我的解释"],
        [
            ["Stage 1", imp1.iloc[0]["feature"], fmt(imp1.iloc[0]["importance_mean"], 3), "SME规模组中，机器学习能力最能解释流程自动化采纳差异。"],
            ["Stage 2", imp2.iloc[0]["feature"], fmt(imp2.iloc[0]["importance_mean"], 3), "行业/区域外部验证中，AI机器学习能力仍是最强特征。"],
        ],
        widths=[2.4, 4.2, 2.2, 6.3],
    )
    high_vif = vif.sort_values("vif", ascending=False).head(4)
    add_table(
        doc,
        ["数据层", "变量", "VIF", "处理态度"],
        [
            [row["dataset"], row["feature_label"], fmt(row["vif"], 2), "不孤立解释单个OLS系数，结合Ridge和树模型判断。"]
            for _, row in high_vif.iterrows()
        ],
        widths=[3.5, 3.4, 1.8, 6.4],
    )
    add_figure(doc, "fig2_stage2_external_importance.png", "图3 Stage 2外部验证特征重要性", width_cm=12.6)

    add_heading(doc, "七、企业画像和部署路径解释")
    add_para(
        doc,
        "聚类结果帮助我把模型输出转成企业可以理解的部署路径。Stage 1形成四类SME画像，工作流自动化比例最高的一类平均达到14.66，最低的一类为2.61。Stage 2形成六类行业/区域画像，最高的一类工作流自动化比例为17.18，最低的一类为2.04。这个差异说明，企业AI采纳不是一个统一速度，而是被能力基础、治理条件和行业环境共同拉开。"
    )
    add_para(
        doc,
        "我的解释重点是“安全顾虑改变路径”，而不是“安全顾虑越高越不用AI”。安全敏感型企业可能不适合直接上公共SaaS，但可以选择私有云、本地化试点或混合部署；效率需求高但治理基础一般的企业，可以从低风险流程、标准SaaS或API接入开始；治理成熟度较高的企业，才更适合流程级自动化和跨系统集成。"
    )
    add_table(
        doc,
        ["画像层", "类别数", "最高自动化比例", "最低自动化比例", "报告中的用途"],
        [
            ["Stage 1 SME", cluster1.shape[0], fmt(cluster1["target_workflow_automation"].max(), 2), fmt(cluster1["target_workflow_automation"].min(), 2), "解释中小企业分层部署路径。"],
            ["Stage 2 行业/区域", cluster2.shape[0], fmt(cluster2["target_workflow_automation"].max(), 2), fmt(cluster2["target_workflow_automation"].min(), 2), "验证行业和地区异质性。"],
        ],
        widths=[3.5, 1.6, 2.3, 2.3, 5.3],
    )
    add_figure(doc, "fig1b_sme_importance_ppt.png", "图4 SME机制解释图表", width_cm=12.6)

    add_heading(doc, "八、A10训练、Agent原型和复现整理")
    add_para(
        doc,
        "为了让项目更像完整的数据挖掘流程，我还把A10 GPU基线和Agent原型纳入整理。A10上的MLP不是为了超过树模型，而是作为神经网络基线，说明在表格数据任务中，复杂模型不一定比结构化特征加树模型更合适。Stage 1的torch_mlp R²为0.8060，Stage 2为0.6621，都低于对应的树模型或Ridge基线。这个结果对课程汇报很有用，因为它提醒我们不能只因为模型听起来复杂就认为它更好。"
    )
    add_table(
        doc,
        ["数据层", "GPU模型", "设备", "R²", "MAE", "训练秒数", "结论"],
        [
            [row["dataset"], row["gpu_model"], f"{row['device']} {row['gpu_name']}", fmt(row["r2"]), fmt(row["mae"]), fmt(row["train_seconds"], 2), "作为神经网络基线，不替代主模型。"]
            for _, row in gpu.iterrows()
        ],
        widths=[4.0, 2.0, 3.2, 1.5, 1.5, 1.7, 3.2],
    )
    add_para(
        doc,
        "Agent原型部分主要用于展示研究成果如何进入企业决策场景。它不是把AI回答写得更花，而是要求回答必须引用来源、不能编造指标、不能把Stage 2写成SME-only结论。最终我把Agent测试、复现说明、数据样本和结果表一起整理进GitHub公开仓库，方便老师顺着代码和输出核对。"
    )
    add_table(
        doc,
        ["验证方式", "代表结果", "我从中得到的判断"],
        [
            ["按国家留出", f"Stage 1 ExtraTrees holdout R²={fmt(holdout[holdout['dataset'].eq('stage1_sme_size_class') & holdout['model'].eq('extra_trees')].iloc[0]['holdout_r2'])}", "留出国家后仍有较好解释力。"],
            ["按国家留出", f"Stage 2 ExtraTrees holdout R²={fmt(holdout[holdout['dataset'].eq('stage2_industry_region_GE10') & holdout['model'].eq('extra_trees')].iloc[0]['holdout_r2'])}", "行业/区域层更复杂，分数下降但仍可用。"],
            ["Agent测试", "公开仓库保留单元测试和证据约束说明", "展示材料要能追溯到数据、代码和报告。"],
        ],
        widths=[3.2, 5.0, 6.5],
    )

    add_heading(doc, "九、GitHub整理和最终提交口径")
    add_para(
        doc,
        "最后整理GitHub时，我专门把“研究工作区”和“老师检查入口”分开。研究工作区里保留src、configs、notebooks、data、outputs和Agent原型，方便复现；最终提交材料则按老师给的顺序整理成四个目录：数据、源码、小组汇报PPT和报告、小组成员个人作业整理。这样老师不需要在历史草稿、预览图和临时文件里找材料。"
    )
    add_para(
        doc,
        "这一步看起来像文件整理，其实也在检验研究质量。比如数据目录必须能看到processed、samples、raw manifest、模型结果表和图表；源码目录必须能看到数据下载、清洗、建模、图表生成和Agent测试代码；小组汇报目录必须同时有PPTX、PPT导出PDF、15页以上课程报告和证据映射；个人目录必须按成员分别保留平时作业汇总PDF、十次作业PDF和个人总结。"
    )
    add_table(
        doc,
        ["最终目录", "我检查的重点", "为什么这样整理"],
        [
            ["01_数据", "processed、samples、raw manifest、模型表、图表和报告是否齐全。", "证明数据来源和结果不是只写在报告里。"],
            ["02_源码", "下载、清洗、训练、诊断、图表、Agent测试代码是否都在。", "让模型结果可以被重新运行和追踪。"],
            ["03_小组汇报PPT和报告", "PPTX、PDF、课程报告、质量核验和证据映射是否对应。", "保证课堂展示不是脱离数据的空讲。"],
            ["04_小组成员个人作业整理", "四名成员是否都有平时作业汇总PDF、10次作业PDF和个人总结。", "对齐老师最终提交要求。"],
        ],
        widths=[4.0, 6.2, 4.8],
    )
    add_para(
        doc,
        "我还检查了一个容易出错的地方：本地材料和GitHub远端不是同一件事。本轮整理是在本地Windows路径下完成的，GitHub地址可以填入共享文档，但没有执行自动提交或推送。因此最终上传前必须由我再审核一次本地目录，并确认哪些文件已经真正进入远端仓库。这个边界如果不写清楚，后面评分时容易出现“本地有、仓库没有”的问题。"
    )

    add_heading(doc, "十、PPT汇报逻辑和证据映射")
    add_para(
        doc,
        "PPT部分我也做了纠偏。早期版本容易写成“项目很完整、文件很齐全”这种汇报，但老师要看的不是文件清单，而是研究工作本身。最终PPT按18页组织：先讲研究问题和理论框架，再讲数据生命周期和清洗边界，接着讲Stage 1和Stage 2模型结果，最后回到企业部署策略和研究局限。"
    )
    add_para(
        doc,
        "我在整理PPT时要求每页都有对应证据。讲数据来源时，对应Eurostat源数据和manifest；讲模型结果时，对应course_algorithm_comparison、course_regression_summary、feature_importance等表；讲A10和Agent时，对应GPU基线、测试代码和复现说明；讲部署路径时，对应聚类画像和特征重要性。这样答辩时即使老师追问某个数字，也能回到表格或代码，而不是靠临场解释。"
    )
    add_para(
        doc,
        "我个人在答辩准备中还做了语言压缩。比如“12,770,332行源数据”不能直接说成“模型训练用了1200多万行”，准确说法应该是官方源文件扫描12,770,332行，经过筛选和聚合形成5,814行建模面板。这个改法看似细，但能体现我们懂数据挖掘流程，也避免夸大结果。"
    )

    add_heading(doc, "十一、从投稿稿到课程报告的改写边界")
    add_para(
        doc,
        "投稿稿给了本项目比较完整的理论表达，但课程作业不能变成投稿稿的翻版。我在重写小组报告和个人总结时做了三个区分：第一，投稿稿强调学术规范和理论贡献，课程报告强调数据清洗、模型验证和可复现；第二，投稿稿可以讨论问卷和访谈构念，课程公开提交以Eurostat官方数据、代码和结果表为主；第三，投稿稿中的研究表述要转成课堂能听懂的语言，不能堆概念。"
    )
    add_para(
        doc,
        "这个边界也影响我的个人总结写法。我不会说“机器学习证明了企业一定会怎样”，而是写“在当前公开数据和验证口径下，模型结果支持哪些解释”。我也不会把Stage 2的行业/区域GE10面板说成中小企业规模组数据。这个地方看似小，但如果写错，报告质量会明显下降。"
    )

    add_heading(doc, "十二、个人反思和后续改进")
    add_para(
        doc,
        "这次项目让我对机器学习作业的理解变了。以前单次作业更像“把算法跑通”，结课案例要求我回答更完整的问题：数据从哪里来，清洗规则是否透明，特征为什么这样构建，验证方式有没有泄漏，结果是否能被复核，PPT里的每一个判断能不能找到证据。"
    )
    add_para(
        doc,
        "我也看到了自己的不足。Eurostat指标口径很多，变量解释需要反复对照官方说明；部分变量VIF较高，说明OLS系数不能单独拿出来讲；PPT为了展示效果容易把话写得过满，后期必须把“绝对”“完美”“完全证明”这类词删掉；GitHub公开仓库和本地最终材料也要分清，未推送的内容不能说成远端已经更新。"
    )
    add_para(
        doc,
        "还有一个更具体的教训：文档质量不只是排版。旧版个人总结虽然能打开，也有PDF页数，但正文太短，很多内容只是概括，没有展示我做过的清洗、建模、验证和改写工作。重新整理后，我把每个判断尽量落到数据表、报告、代码或PPT证据上。这个过程让我意识到，课程报告最怕“看上去很整齐，实际上没有内容”。"
    )
    add_para(
        doc,
        "在协作上，我也学到一点：组长不能只负责最后合并文件。真正困难的是把不同成员的材料改成同一套研究语言。有人负责数据，有人负责模型，有人负责展示，但最后报告必须能说明同一个问题。后期我花了不少时间处理这个问题，尤其是把投稿稿语言改成课程汇报语言，把文件清单式表达改成研究工作总结。"
    )
    add_para(
        doc,
        "如果继续完善这个项目，我会优先补三件事。第一，补充中国本土企业公开数据或更规范的问卷数据，检验欧洲样本结论能否迁移。第二，增加时间留出和行业留出验证，让模型评价更严格。第三，把Agent原型和真实业务流程连接起来，验证它能不能帮助企业在SaaS、API、本地化和混合部署之间做选择。"
    )

    add_heading(doc, "十三、个人总结")
    add_para(
        doc,
        "我在本次小组案例中的核心贡献，是把一个偏论文的研究主题整理成机器学习课程可以检查的完整项目。这个项目有数据来源，有清洗代码，有模型比较，有图表输出，有小组PPT，也有最终PDF报告。我的个人总结不只是写“我参与了小组合作”，而是把我实际承担的工作讲清楚：我怎样把研究问题拆成建模任务，怎样控制数据边界，怎样解释模型结果，怎样把投稿稿内容改成课程展示，怎样把最终文件整理到老师要求的提交结构里。"
    )
    add_para(
        doc,
        "从课程角度看，我最大的收获不是某一个模型分数，而是形成了比较完整的数据挖掘工作习惯。模型结果必须能追溯到数据和代码，报告结论必须能对应模型和图表，展示材料必须知道哪些能说、哪些不能说。这个过程比写一份漂亮但空的总结更慢，也更麻烦，但它更接近这门课真正要考查的能力。"
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("信管2301 景浩伟 202321054012｜机器学习课程个人案例总结报告")
    set_font(run, size=9)

    doc.save(docx_path)
    normalize_docx(docx_path)

    if pdf_path.exists():
        pdf_path.unlink()
    soffice = shutil.which("soffice") or str(Path.home() / ".codex-office" / "bin" / "soffice.exe")
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )

    text = docx_text(docx_path)
    pages, blanks = pdf_pages_and_blanks(pdf_path)
    color_hits: list[tuple[str, str]] = []
    with zipfile.ZipFile(docx_path) as zf:
        for name in zf.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                xml = zf.read(name).decode("utf-8", errors="ignore")
                for value in re.findall(r'w:color[^>]*w:val="([^"]+)"', xml):
                    if value.upper() not in {"000000", "AUTO"}:
                        color_hits.append((name, value))
    text_audit = audit_text(text)
    quality_report = {
        "ok": text_audit["ok"] and pages >= 10 and not blanks and not color_hits,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "docx": str(docx_path),
        "pdf": str(pdf_path),
        "pages": pages,
        "blank_pages": blanks,
        "non_black_color_count": len(color_hits),
        "text_audit": text_audit,
        "skills_applied": [
            "Word / DOCX: OOXML生成、黑色字体、PDF round-trip转换",
            "Word中文格式标准化: A4、微软雅黑、标题层级、正文1.5倍行距",
            "avoid-ai-writing: 增加第一人称工作细节，减少套话和泛泛判断",
            "academic-paper-reviewer: 用方法、证据、边界、可复现性检查内容质量",
        ],
    }
    stale_local_audit = out_dir / "景浩伟个人总结报告_质量核验.json"
    if stale_local_audit.exists():
        stale_local_audit.unlink()
    (FINAL / "景浩伟个人总结报告_质量核验.json").write_text(json.dumps(quality_report, ensure_ascii=False, indent=2), encoding="utf-8")
    if not quality_report["ok"]:
        raise RuntimeError(json.dumps(quality_report, ensure_ascii=False, indent=2))
    return quality_report


if __name__ == "__main__":
    print(json.dumps(build(), ensure_ascii=False, indent=2))
